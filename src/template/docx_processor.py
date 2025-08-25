"""
DOCX template processing module for CVPilot
Handles direct text replacement while preserving document styles
"""

import shutil
import re
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple

from docx import Document
from docx.shared import Inches
from docx.oxml.shared import qn
from docx.oxml import OxmlElement

from ..utils.models import Replacements
from ..utils.logger import LoggerMixin
from ..utils.naming_utils import NamingUtils
from ..utils.business_rules_engine import BusinessRulesEngine

class DocxProcessor(LoggerMixin):
    """Process DOCX templates with direct text replacement"""
    
    def __init__(self):
        super().__init__()

        # Initialize Business Rules Engine for validation
        self.business_rules = BusinessRulesEngine()

        # Define the specific sections to replace based on actual template structure
        self.target_sections = {
            'cv_title': ['PRODUCT ANALYST', 'CURRICULUM VITAE', 'CV', 'RESUME', 'CURRÍCULUM VITAE', 'INNOVATION SPECIALIST', 'PROJECT MANAGER', 'DATA ANALYST', 'BUSINESS ANALYST', 'PRODUCT OWNER', 'OPERATIONS MANAGER'],
            'professional_summary': ['PRODUCT MANAGER WITH', 'PROFESSIONAL SUMMARY', 'RESUMEN PROFESIONAL', 'SUMMARY', 'EXECUTIVE SUMMARY', 'RESULTS-ORIENTED', 'EXPERIENCED', 'SENIOR'],
            'skills_software': ['SKILLS', 'SKILLS & SOFTWARE', 'SKILLS AND SOFTWARE', 'TECHNICAL SKILLS', 'COMPETENCIAS'],
            'experience_titles': ['PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'EXPERIENCIA PROFESIONAL', 'WORK EXPERIENCE', 'EMPLOYMENT HISTORY']
        }
    
    def create_backup(self, template_path: Path, backups_path: Path) -> Path:
        """Create backup of template file"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"template_backup_{timestamp}.docx"
        backup_path = backups_path / backup_name
        
        try:
            shutil.copy2(template_path, backup_path)
            self.log_info(f"Template backup created: {backup_path}")
            return backup_path
        except Exception as e:
            self.log_error(f"Failed to create backup: {str(e)}")
            raise
    
    def process_template(self, template_path: Path, replacements: Replacements, output_path: Path, 
                        job_category: str = "general") -> Path:
        """Process template with replacements and generate final CV"""
        
        self.log_info(f"Processing template: {template_path}")
        
        try:
            # Create backup first
            backups_path = output_path.parent / "backups"
            backups_path.mkdir(exist_ok=True)
            self.create_backup(template_path, backups_path)
            
            # Load document
            doc = Document(str(template_path))
            
            # Apply replacements
            self._apply_replacements(doc, replacements)
            
            # Generate intelligent folder name
            software_list = replacements.software_list.content.split(', ') if isinstance(replacements.software_list.content, str) else replacements.software_list.content
            folder_name = NamingUtils.generate_folder_name(replacements.position, software_list)
            category_path = output_path / folder_name
            category_path.mkdir(exist_ok=True)
            
            # Generate intelligent filename
            output_filename = NamingUtils.generate_filename(replacements.position, software_list, replacements.company)
            output_file = category_path / output_filename
            
            # Save processed document - NO OVERWRITE RULE
            self._save_without_overwrite(doc, output_file)
            
            self.log_info(f"CV generated successfully: {output_file}")
            self.log_info(f"Folder: {folder_name}")
            return output_file
            
        except Exception as e:
            self.log_error(f"Template processing failed: {str(e)}")
            raise
    
    def generate_dry_run(self, template_path: Path, replacements: Replacements, output_path: Path) -> Path:
        """Generate dry-run version (DOCX preview) for review"""
        
        self.log_info("Generating dry-run preview")
        
        try:
            # Load document
            doc = Document(str(template_path))
            
            # Apply replacements
            self._apply_replacements(doc, replacements)
            
            # Save as preview DOCX with intelligent naming - NO OVERWRITE RULE
            software_list = replacements.software_list.content.split(', ') if isinstance(replacements.software_list.content, str) else replacements.software_list.content
            preview_filename = NamingUtils.generate_filename(replacements.position, software_list, replacements.company)
            preview_filename = f"Preview_{preview_filename}"
            preview_file = output_path / preview_filename

            self._save_without_overwrite(doc, preview_file)
            
            self.log_info(f"Dry-run preview generated: {preview_file}")
            return preview_file
            
        except Exception as e:
            self.log_error(f"Dry-run generation failed: {str(e)}")
            raise
    
    def _apply_replacements(self, doc: Document, replacements: Replacements):
        """Apply all replacements to the document while preserving styles"""
        
        self.log_info("🚀 Starting document replacements...")
        
        # Track replacement results
        replacement_results = {
            'cv_title': False,
            'professional_summary': False,
            'skills_software': False,
            'experience_titles': False
        }
        
        # 1. Replace CV title (main role)
        self.log_info("=" * 60)
        self.log_info("📝 STEP 1: CV TITLE REPLACEMENT")
        self.log_info("=" * 60)
        replacement_results['cv_title'] = self._replace_cv_title(doc, replacements.objective_title.content)
        
        # 2. Replace professional summary
        self.log_info("=" * 60)
        self.log_info("📝 STEP 2: PROFESSIONAL SUMMARY REPLACEMENT")
        self.log_info("=" * 60)
        replacement_results['professional_summary'] = self._replace_professional_summary(doc, replacements.profile_summary.content)
        
        # 3. Replace skills and software (two separate lines)
        self.log_info("=" * 60)
        self.log_info("📝 STEP 3: SKILLS & SOFTWARE REPLACEMENT")
        self.log_info("=" * 60)
        replacement_results['skills_software'] = self._replace_skills_software(doc, replacements.skill_list.content, replacements.software_list.content)
        
        # 4. Replace experience job titles
        self.log_info("=" * 60)
        self.log_info("📝 STEP 4: EXPERIENCE TITLES REPLACEMENT")
        self.log_info("=" * 60)
        replacement_results['experience_titles'] = self._replace_experience_titles(doc, replacements.objective_title.content)
        
        # Summary of replacements
        self.log_info("=" * 60)
        self.log_info("📊 REPLACEMENT SUMMARY")
        self.log_info("=" * 60)
        
        successful_replacements = sum(replacement_results.values())
        total_replacements = len(replacement_results)
        
        for section, success in replacement_results.items():
            status = "✅ SUCCESS" if success else "❌ FAILED"
            self.log_info(f"{section.replace('_', ' ').title()}: {status}")
        
        self.log_info(f"Overall: {successful_replacements}/{total_replacements} replacements successful")
        
        if successful_replacements == total_replacements:
            self.log_info("🎉 All replacements completed successfully!")
        else:
            self.log_warning(f"⚠️ {total_replacements - successful_replacements} replacements failed")
        
        return replacement_results
    
    def _replace_cv_title(self, doc: Document, new_title: str):
        """Replace the main CV title with the job target title - VALIDATED by business rules"""
        # Convert to uppercase for consistency
        new_title_upper = new_title.upper()
        self.log_info(f"🔍 Looking for CV title to replace with: '{new_title_upper}'")
        self.log_info(f"🔍 DEBUG: Searching in {len(doc.paragraphs)} paragraphs")
        self.log_info(f"🔍 DEBUG: Target keywords: {self.target_sections['cv_title']}")

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                self.log_info(f"🔍 DEBUG: Paragraph {i}: '{text}' (length: {len(text)})")
                text_upper = text.upper()

                # Check each keyword individually
                for keyword in self.target_sections['cv_title']:
                    if keyword in text_upper:
                        self.log_info(f"🔍 DEBUG: FOUND MATCH - Keyword '{keyword}' found in paragraph {i}")

                        original_text = text

                        # USE BUSINESS RULES ENGINE FOR VALIDATION
                        self.log_info(f"🔍 DEBUG: Calling validate_cv_title_replacement with:")
                        self.log_info(f"   current_title: '{original_text}'")
                        self.log_info(f"   target_title: '{new_title}'")

                        validation = self.business_rules.validate_cv_title_replacement(
                            current_title=original_text,
                            target_title=new_title
                        )

                        self.log_info(f"🔍 DEBUG: Validation result: {validation}")

                        if not validation['should_replace']:
                            self.log_info(f"✅ CV TITLE VALIDATION: {validation['explanation']}")
                            return False  # Don't replace if validation says no

                        # Replace if validation allows it
                        original_style = paragraph.style
                        paragraph.text = new_title_upper
                        paragraph.style = original_style
                        self.log_info(f"✅ CV TITLE REPLACED: '{original_text}' → '{new_title_upper}'")
                        return True

        self.log_warning(f"⚠️ CV title section not found. Keywords searched: {self.target_sections['cv_title']}")
        return False

    def audit_replacements(self, replacements_log: List[Dict]) -> Dict[str, Any]:
        """Audit replacements using the business rules engine"""
        return self.business_rules.audit_replacements(replacements_log)
    
    def _replace_professional_summary(self, doc: Document, new_summary: str):
        """Replace the professional summary section"""
        self.log_info("🔍 Looking for professional summary section to replace")
        
        # Look for the paragraph that contains the professional summary
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Check if this paragraph contains the professional summary
            if any(keyword in text.upper() for keyword in self.target_sections['professional_summary']):
                # This is the summary paragraph itself
                original_style = paragraph.style
                original_text = text
                original_length = len(original_text)
                
                # Use the summary as generated (no truncation needed)
                new_summary_adjusted = new_summary
                
                paragraph.text = new_summary_adjusted
                paragraph.style = original_style
                
                self.log_info(f"✅ PROFESSIONAL SUMMARY REPLACED: '{original_text[:100]}...' → '{new_summary_adjusted[:100]}...'")
                self.log_info(f"📏 Length: {original_length} → {len(new_summary_adjusted)} chars (no truncation)")
                return True
        
        self.log_warning(f"⚠️ Professional summary section not found")
        return False
    
    def _replace_skills_software(self, doc: Document, skills_content: str, software_content: str):
        """Replace the skills and software section (two separate lines)"""
        self.log_info("🔍 Looking for skills and software section to replace")
        
        # Find the skills section and replace the two paragraphs after the header
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Check if this is the skills section header (more specific detection)
            if any(keyword in text.upper() for keyword in self.target_sections['skills_software']):
                # Additional check: make sure this is not the professional summary
                if len(text) < 50 and not any(word in text.lower() for word in ['experience', 'years', 'proven', 'ability']):
                    self.log_info(f"📍 Found skills section header: '{text}'")
                    
                    # Replace the next two paragraphs (skills and software content)
                    replacements_made = 0
                    
                    # First paragraph (skills)
                    if i + 1 < len(doc.paragraphs):
                        skills_paragraph = doc.paragraphs[i + 1]
                        if skills_paragraph.text.strip():  # Make sure it has content
                            original_style = skills_paragraph.style
                            original_text = skills_paragraph.text
                            original_length = len(original_text)
                            
                            # Adjust skills content length to match original (±10 characters)
                            target_length = original_length
                            if len(skills_content) > target_length + 10:
                                # Truncate if too long - find a good breaking point
                                max_length = target_length + 10
                                words = skills_content.split()
                                truncated_text = ""
                                for word in words:
                                    if len(truncated_text + " " + word) <= max_length:
                                        truncated_text += (" " + word) if truncated_text else word
                                    else:
                                        break
                                skills_adjusted = truncated_text
                            elif len(skills_content) < target_length - 10:
                                # Pad if too short
                                padding_needed = target_length - len(skills_content) - 10
                                skills_adjusted = skills_content + " " * padding_needed
                            else:
                                skills_adjusted = skills_content
                            
                            skills_paragraph.text = skills_adjusted
                            skills_paragraph.style = original_style
                            
                            self.log_info(f"✅ SKILLS LINE REPLACED: '{original_text[:100]}...' → '{skills_adjusted[:100]}...'")
                            self.log_info(f"📏 Skills Length: {original_length} → {len(skills_adjusted)} chars (target: {target_length}±10)")
                            replacements_made += 1
                    
                    # Second paragraph (software)
                    if i + 2 < len(doc.paragraphs):
                        software_paragraph = doc.paragraphs[i + 2]
                        if software_paragraph.text.strip():  # Make sure it has content
                            original_style = software_paragraph.style
                            original_text = software_paragraph.text
                            original_length = len(original_text)
                            
                            # Adjust software content length to match original (±10 characters)
                            target_length = original_length
                            if len(software_content) > target_length + 10:
                                # Truncate if too long - find a good breaking point
                                max_length = target_length + 10
                                words = software_content.split()
                                truncated_text = ""
                                for word in words:
                                    if len(truncated_text + " " + word) <= max_length:
                                        truncated_text += (" " + word) if truncated_text else word
                                    else:
                                        break
                                software_adjusted = truncated_text
                            elif len(software_content) < target_length - 10:
                                # Pad if too short
                                padding_needed = target_length - len(software_content) - 10
                                software_adjusted = software_content + " " * padding_needed
                            else:
                                software_adjusted = software_content
                            
                            software_paragraph.text = software_adjusted
                            software_paragraph.style = original_style
                            
                            self.log_info(f"✅ SOFTWARE LINE REPLACED: '{original_text[:100]}...' → '{software_adjusted[:100]}...'")
                            self.log_info(f"📏 Software Length: {original_length} → {len(software_adjusted)} chars (target: {target_length}±10)")
                            replacements_made += 1
                    
                    if replacements_made > 0:
                        self.log_info(f"✅ Replaced {replacements_made} skills/software lines")
                        return True
        
        self.log_warning(f"⚠️ Skills and software content not found")
        return False
    
    def _replace_experience_titles(self, doc: Document, new_title: str):
        """Replace job titles in experience section"""
        self.log_info(f"🔍 Looking for experience job titles to replace with: '{new_title}'")
        
        in_experience_section = False
        replacements_made = 0
        experience_entries = []
        paragraphs_to_remove = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Check if we're entering the experience section
            if any(keyword in text.upper() for keyword in self.target_sections['experience_titles']):
                in_experience_section = True
                self.log_info(f"📍 Found experience section header: '{text}'")
                continue
            
            # If we're in experience section, look for job titles or company names
            if in_experience_section and text:
                # Look for patterns that indicate job titles OR company names in experience section
                if self._is_experience_job_title(text):
                    # Check if this title should be replaced (only GCA advanced profile titles)
                    if self._should_replace_title(text, doc, new_title):
                        # Preserve original style
                        original_style = paragraph.style
                        original_text = text
                        
                        # Extract company name and location from original text
                        company_name = self._extract_company_name(original_text)
                        location = self._extract_location(original_text)
                        
                        # Create new job title with proper formatting and capitalization
                        # For job titles, we want to replace the job title part, not add a prefix
                        # Ensure proper capitalization
                        capitalized_title = new_title.title() if new_title else new_title
                        
                        # Extract date and specialization from original text (preserve original formatting)
                        specialization = ""
                        date_part = ""

                        # Extract specialization from original text (preserve original formatting)
                        if '(' in original_text and ')' in original_text:
                            start = original_text.find('(')
                            end = original_text.find(')')
                            if start != -1 and end != -1 and end > start:
                                specialization = original_text[start:end+1]  # Include parentheses

                        # Extract date part (support multiple formats)
                        import re
                        date_patterns = [
                            r'(\d{2}/\d{4})\s*[-–]\s*(\d{2}/\d{4})',  # 08/2022 - 11/2023
                            r'(\d{2}/\d{4})\s*[-–]\s*Present',       # 11/2023 - Present
                            r'(\d{2}/\d{4})\s*Present',              # 11/2023 Present (sin guion)
                            r'(\d{2}/\d{4})\s+(\d{2}/\d{4})',        # 08/2022 11/2023 (sin guion)
                        ]

                        for pattern in date_patterns:
                            date_match = re.search(pattern, original_text)
                            if date_match:
                                if len(date_match.groups()) == 2:
                                    date_part = f"{date_match.group(1)} - {date_match.group(2)}"
                                else:
                                    date_part = f"{date_match.group(1)} - Present"
                                break
                        
                        # BULLET POOL RULE: Simple EXACT replacement
                        # Replace the entire original text with the target job title
                        # NO text parsing, NO format changes, NO date manipulation

                        new_job_title = self._selected_title
                        
                        # Replace the text while preserving the original font style AND dates
                        original_runs = list(paragraph.runs)
                        if original_runs:
                            # Extract the job title part (usually the first part before dates)
                            original_text = paragraph.text

                            # Look for date patterns in the original text
                            date_patterns = [
                                r'(\d{1,2}/\d{4})\s*[-–—]\s*(\d{1,2}/\d{4})',  # MM/YYYY - MM/YYYY
                                r'(\d{1,2}/\d{4})\s*[-–—]\s*Present',         # MM/YYYY - Present
                                r'(\d{1,2}/\d{4})\s*[-–—]\s*Current',         # MM/YYYY - Current
                                r'(\d{4})\s*[-–—]\s*(\d{4})',                 # YYYY - YYYY
                                r'(\d{4})\s*[-–—]\s*Present',                 # YYYY - Present
                                r'(\d{4})\s*[-–—]\s*Current',                 # YYYY - Current
                            ]

                            dates_found = None
                            for pattern in date_patterns:
                                match = re.search(pattern, original_text)
                                if match:
                                    dates_found = match.group(0)
                                    break

                            # Replace the entire text but preserve dates if found
                            if dates_found:
                                # Keep the date part and replace only the title
                                new_text = f"{new_job_title} {dates_found}"
                                paragraph.text = new_text
                                self.log_info(f"✅ EXPERIENCE TITLE REPLACED WITH DATES: '{original_text}' → '{new_text}'")
                            else:
                                # Fallback: replace just the first run if no dates found
                                first_run = original_runs[0]
                                first_run.text = new_job_title

                                # Clear any additional runs that don't contain dates
                                for run in original_runs[1:]:
                                    run_text = run.text.strip()
                                    if not any(char.isdigit() or char in ['-', '–', '—', '/', 'Present', 'Current'] for char in run_text):
                                        run.text = ""
                        else:
                            # If no runs exist, just set the text
                            paragraph.text = new_job_title
                        paragraph.style = original_style
                        replacements_made += 1
                        experience_entries.append({
                            'original': original_text,
                            'new': new_job_title,
                            'company': company_name,
                            'location': location
                        })
                        self.log_info(f"✅ EXPERIENCE TITLE REPLACED: '{original_text}' → '{new_job_title}'")
                    else:
                        self.log_info(f"🚫 SKIPPING TITLE: '{text}' (no alternatives in bullet pool)")
                
                # Also look for and remove duplicate job titles that might appear right after
                # These are typically the original job titles that should be removed
                elif self._is_duplicate_job_title(text, new_title):
                    paragraphs_to_remove.append(i)
                    self.log_info(f"🗑️ MARKED FOR REMOVAL: '{text}' (duplicate job title)")
        
        # Remove duplicate paragraphs (in reverse order to maintain indices)
        for i in reversed(paragraphs_to_remove):
            if i < len(doc.paragraphs):
                removed_text = doc.paragraphs[i].text.strip()
                doc.paragraphs[i].text = ""  # Clear the paragraph
                self.log_info(f"🗑️ REMOVED DUPLICATE: '{removed_text}'")
        
        if replacements_made > 0:
            self.log_info(f"✅ Replaced {replacements_made} experience job titles")
            self.log_info(f"📊 Experience entries updated:")
            for entry in experience_entries:
                self.log_info(f"   - {entry['company']}: {entry['original']} → {entry['new']}")
        else:
            self.log_warning(f"⚠️ No experience job titles found to replace")
            self.log_info(f"💡 This might be normal if the template doesn't have experience entries yet")
        
        return replacements_made > 0
    
    def _is_experience_job_title(self, text: str) -> bool:
        """Determine if a paragraph contains an experience job title"""
        # Experience job titles are usually:
        # - Contain ONLY job title keywords (Analyst, Manager, Specialist, etc.)
        # - Have a specific format: "Job Title (Application/Platform)" or "Job Title"
        # - Are in Normal style (not List Paragraph)
        # - Don't start with bullet points
        # - Usually followed by dates
        # - Do NOT contain company names or locations
        
        if not text or text.startswith('•') or text.startswith('-') or text.startswith('*'):
            return False
        
        # Look for job title patterns that indicate this is a job title
        job_title_patterns = [
            'Product Analyst', 'Product Operations Specialist', 'Quality Assurance Analyst',
            'Product Manager', 'Project Manager', 'Business Analyst', 'Data Analyst',
            'Senior', 'Junior', 'Lead', 'Principal', 'Head', 'Director',
            'Analyst', 'Manager', 'Specialist', 'Coordinator', 'Developer'
        ]
        
        # Must contain at least one job title pattern
        has_job_title = any(pattern in text for pattern in job_title_patterns)
        
        # Must contain application/platform indicators or parentheses
        has_application = any(indicator in text for indicator in ['(', ')', 'Application', 'Platform', 'SaaS', 'B2B', 'B2C'])
        
        # Must have date indicators (usually at the end)
        has_dates = any(indicator in text for indicator in ['Present', '2023', '2022', '2021', '2020', '2019', '–', '—'])
        
        # Should be reasonably short (not a full paragraph)
        is_reasonable_length = len(text) < 100
        
        # Should not contain typical paragraph content
        not_paragraph_content = not any([
            text.endswith('.'),
            len(text.split()) > 15,
            'with' in text.lower() and len(text) > 80,
            'experience' in text.lower() and len(text) > 80,
            'responsibilities' in text.lower(),
            'duties' in text.lower(),
            'managed' in text.lower() and len(text) > 60
        ])
        
        # CRITICAL: Should NOT contain company names or locations
        company_indicators = [
            'Growing Companies Advisors', 'GCA', 'Industrias de Tapas Taime', 'Loszen', 
            'Industrias QProductos', 'Consulting firm', 'Manufacturing company', 'startup',
            'U.S.', 'Venezuela', 'Spain', 'Remote', 'Caracas', 'Santa Cruz de Aragua'
        ]
        
        has_company_info = any(indicator in text for indicator in company_indicators)
        
        # Additional check: should look like a job title format
        looks_like_title = (
            has_job_title and 
            (has_application or has_dates) and 
            is_reasonable_length and 
            not_paragraph_content and
            not has_company_info  # CRITICAL: No company information
        )
        
        return looks_like_title

    def _is_company_in_experience(self, text: str) -> bool:
        """Check if text looks like a company name in experience section"""
        # Company patterns in the template we're using:
        # - Contains company indicators (commas, parentheses, em-dashes, C.A, S.A)
        # - Contains company names or location information
        # - Is in the experience section
        # - Is not a bullet point or description

        if not text or text.startswith('•') or text.startswith('-') or text.startswith('*'):
            return False

        # Company indicators
        has_company_indicators = any(char in text for char in [',', '(', ')', '—', '-', 'C.A', 'S.A'])

        # Company names and locations
        company_names = [
            'Growing Companies Advisors', 'GCA', 'Industrias de Tapas Taime',
            'Loszen', 'Industrias QProductos', 'Consulting firm',
            'Manufacturing company', 'startup', 'U.S.', 'Venezuela', 'Spain',
            'Remote', 'Caracas', 'Santa Cruz de Aragua'
        ]

        contains_company_info = any(name.lower() in text.lower() for name in company_names)

        # Should be reasonably short (company names with locations)
        is_reasonable_length = 10 < len(text) < 150

        # Should not be a long description paragraph
        not_long_description = len(text.split()) < 20

        return (has_company_indicators or contains_company_info) and is_reasonable_length and not_long_description

    def _analyze_responsibilities_for_role(self, company_text: str, responsibilities: list) -> str:
        """Analyze responsibilities text to determine the most appropriate role title"""
        all_text = ' '.join(responsibilities).lower()

        # Keywords for different roles
        role_keywords = {
            'Product Manager': ['product manager', 'product strategy', 'roadmap', 'stakeholder', 'backlog', 'sprint', 'user story', 'requirements', 'prioritization'],
            'Project Manager': ['project manager', 'project management', 'timeline', 'milestone', 'deliverable', 'coordination', 'planning', 'execution', 'deadline'],
            'Business Analyst': ['business analyst', 'requirements gathering', 'business process', 'workflow', 'documentation', 'analysis', 'reporting'],
            'Product Owner': ['product owner', 'agile', 'scrum', 'user story', 'acceptance criteria', 'backlog refinement', 'sprint planning'],
            'Product Analyst': ['product analyst', 'metrics', 'kpi', 'analytics', 'data analysis', 'reporting', 'insights', 'performance'],
            'Quality Assurance Analyst': ['quality assurance', 'testing', 'qa', 'test case', 'defect', 'bug', 'validation', 'verification'],
            'Digital Product Specialist': ['digital product', 'saas', 'platform', 'user experience', 'ux', 'interface', 'design', 'development']
        }

        # Score each role based on keyword matches
        role_scores = {}
        for role, keywords in role_keywords.items():
            score = sum(1 for keyword in keywords if keyword in all_text)
            role_scores[role] = score

        # Get the role with the highest score
        if role_scores:
            best_role = max(role_scores, key=role_scores.get)
            if role_scores[best_role] > 0:  # At least one keyword match
                return best_role

        # Default fallback based on company
        if 'gca' in company_text.lower():
            return "Product Manager"
        elif 'loszen' in company_text.lower():
            return "Product Manager"
        else:
            return "Business Analyst"  # Default for manufacturing companies

    def _extract_company_name(self, text: str) -> str:
        """Extract company name from job title text"""
        # For job titles, we want to extract application/platform info, not company names
        # Common patterns:
        # "Job Title (Application/Platform)"
        # "Job Title"
        
        # Look for parentheses which usually contain application/platform info
        if '(' in text and ')' in text:
            # Extract the part inside parentheses
            start = text.find('(')
            end = text.find(')')
            if start != -1 and end != -1 and end > start:
                app_part = text[start+1:end].strip()
                # Only return if it looks like application/platform info, not company name
                if any(keyword in app_part for keyword in ['Application', 'Platform', 'SaaS', 'B2B', 'B2C', 'System', 'essentials']):
                    return app_part
        
        # If no parentheses, try to extract from the end (after job title)
        # Look for common application/platform keywords
        application_keywords = ['Application', 'Platform', 'SaaS', 'B2B', 'B2C', 'System']
        for keyword in application_keywords:
            if keyword in text:
                # Find the position of the keyword and extract from there
                pos = text.find(keyword)
                if pos > 0:
                    # Look backwards for the start of the application name
                    # Usually it's after a space or special character
                    for i in range(pos-1, -1, -1):
                        if text[i] in [' ', '\t', '(', '-']:
                            app_part = text[i+1:].strip()
                            return app_part
        
        # If no clear pattern found, return empty string (no company/app info)
        return ""
    
    def _extract_location(self, text: str) -> str:
        """Extract location from job title text"""
        # For job titles, location is usually not specified
        # Return empty string as job titles typically don't include location
        return ""
    
    def _is_job_title(self, text: str) -> bool:
        """Determine if a paragraph contains a job title"""
        # Job titles are usually:
        # - Short (less than 100 characters)
        # - Contain common job title keywords
        # - Don't contain typical paragraph content
        
        if len(text) > 100:
            return False
        
        job_keywords = [
            'manager', 'director', 'analyst', 'specialist', 'coordinator', 'lead',
            'senior', 'junior', 'principal', 'head', 'chief', 'vice president',
            'product', 'marketing', 'sales', 'engineering', 'design', 'data',
            'business', 'strategy', 'operations', 'customer', 'user', 'growth'
        ]
        
        text_lower = text.lower()
        
        # Check if it contains job keywords
        has_job_keywords = any(keyword in text_lower for keyword in job_keywords)
        
        # Check if it doesn't look like regular paragraph content
        is_not_paragraph = not any([
            text.startswith('•'), text.startswith('-'), text.startswith('*'),
            text.endswith('.'), len(text.split()) > 15
        ])
        
        return has_job_keywords and is_not_paragraph
    
    def _should_replace_title(self, text: str, doc: Document, job_target: str = None) -> bool:
        """Determine if a title should be replaced based on bullet pool alternatives"""
        # Only replace titles that are in GCA advanced profile tables
        # These are the only titles that have alternatives in the bullet pool
        
        # Extract the job title part (before parentheses)
        # Extract the job title part (before parentheses, dates, and extra spaces)
        # Remove dates and clean the title
        import re
        # Remove date patterns and clean up
        title_without_dates = re.sub(r'\d{2}/\d{4}', '', text)
        title_without_dates = re.sub(r'Present', '', title_without_dates)
        title_without_dates = re.sub(r'\s+', ' ', title_without_dates)  # Remove extra spaces
        title_without_dates = re.sub(r'[-–—]+', '', title_without_dates)  # Remove dashes
        job_title_part = title_without_dates.split('(')[0].strip().lower()
        
        # STRICT RULES based on bullet_pool.docx tables:
        # ONLY titles that exist EXACTLY in the bullet_pool.docx tables can be replaced
        # No creative titles or modifications - stick to what's in the pool
        replaceable_titles = [
            'product manager',  # From Table 6 and Table 9
            'product owner',    # From Table 6 and Table 9
            'product analyst',  # From Table 6
            'business analyst', # From Table 6 and Table 9
            'project manager',  # From Table 6 and Table 9 (FALTABA!)
            'product operations specialist',  # From Table 7 (FALTABA!)
            'quality assurance analyst',  # From Table 8
            'quality analyst'  # From Table 3
        ]
        
        # Debug logging
        self.log_info(f"🔍 Checking title: '{job_title_part}' against replaceable_titles: {replaceable_titles}")
        self.log_info(f"🔍 Is in replaceable_titles: {job_title_part in replaceable_titles}")
        
        # CRITICAL FIX: If current title IS in bullet pool, DO NOT REPLACE
        # This is the fundamental rule that was broken
        if job_title_part in replaceable_titles:
            should_replace = False  # Already valid, no replacement needed
            self.log_info(f"   📋 Title '{job_title_part}' already in bullet pool: should_replace = False")
        else:
            should_replace = True  # Not in bullet pool, can potentially replace
        
        # STRICT RULES from bullet_pool.docx:
        # Only replace GCA advanced profile titles with EXACT matches
        
        # Check if this is a GCA company entry (look in previous lines for context)
        is_gca_context = False
        is_loszen_context = False
        
        # Look for company context in previous lines
        current_line_index = None
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip() == text.strip():
                current_line_index = idx
                break
        
        if current_line_index is not None:
            self.log_info(f"🔍 Checking context for title at line {current_line_index}: '{text}'")
            for prev_line in range(max(0, current_line_index-3), current_line_index):
                if prev_line < len(doc.paragraphs):
                    prev_text = doc.paragraphs[prev_line].text.lower()
                    self.log_info(f"   📋 Line {prev_line}: '{prev_text}'")
                    if 'gca' in prev_text or 'growing companies' in prev_text:
                        is_gca_context = True
                        self.log_info(f"   ✅ GCA context found in line {prev_line}")
                    elif 'loszen' in prev_text:
                        is_loszen_context = True
                        self.log_info(f"   ✅ Loszen context found in line {prev_line}")
        else:
            self.log_info(f"🚫 Could not find line index for text: '{text}'")
        
        # USE BUSINESS RULES ENGINE FOR VALIDATION
        context_company = "GCA" if is_gca_context else ("Loszen" if is_loszen_context else None)

        validation = self.business_rules.validate_title_replacement(
            current_title=text,
            target_title=job_target,
            context_company=context_company
        )

        should_replace = validation['should_replace']
        
        # BULLET POOL RULE: Use the target job title EXACTLY as provided
        # NO modifications, NO intelligence, NO creative replacements
        if should_replace and job_target:
            self._selected_title = job_target.strip()
        else:
            self._selected_title = None
        if should_replace:
            self.log_info(f"✅ Title '{text}' should be replaced (found in bullet pool rules)")
        else:
            self.log_info(f"🚫 Title '{text}' should NOT be replaced (not in bullet pool rules)")
        
        return should_replace
    

    
    def _select_intelligent_title(self, job_target: str, current_title: str, is_gca_context: bool, is_loszen_context: bool) -> str:
        """CRITICAL FIX: Use job target directly - NO bullet pool selection"""
        # BULLET POOL RULE VIOLATION FIX:
        # This function was incorrectly selecting from bullet pool options
        # It should use the JOB TARGET directly

        return job_target.strip()
    
    def _extract_period_from_title(self, title_text: str) -> str:
        """Extract period from title text"""
        # Look for date patterns
        import re
        date_pattern = r'(?:Present|\d{2}/\d{4})'
        dates = re.findall(date_pattern, title_text)
        
        if len(dates) >= 2:
            start_date = dates[0]
            end_date = dates[1]
            return f"{start_date}-{end_date}"
        elif len(dates) == 1:
            if dates[0] == "Present":
                return "11/2023-Present"  # Assume current period
            else:
                return f"{dates[0]}-Present"  # Assume current period
        
        return "Unknown"
    
    def _is_duplicate_job_title(self, text: str, new_title: str) -> bool:
        """Determine if a paragraph contains a duplicate job title that should be removed"""
        # BULLET POOL RULE: Only remove duplicates if they contain the EXACT new title
        # NO creative pattern matching, NO guessing
        
        if not text or text.startswith('•') or text.startswith('-') or text.startswith('*'):
            return False
        
        # If this text contains the EXACT new title, it's a duplicate to remove
        return new_title.strip().lower() in text.lower()
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe file system usage"""
        # Remove or replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Limit length
        if len(filename) > 50:
            filename = filename[:50]
        
        return filename.strip()
    
    def validate_template(self, template_path: Path) -> Dict[str, Any]:
        """Validate template file and check for required sections"""
        
        validation_result = {
            "is_valid": True,
            "errors": [],
            "warnings": [],
            "sections_found": [],
            "missing_sections": []
        }
        
        try:
            # Load document
            doc = Document(str(template_path))
            
            # Extract text content
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Check for required sections
            found_sections = []
            for section_name, keywords in self.target_sections.items():
                section_found = any(keyword in full_text.upper() for keyword in keywords)
                if section_found:
                    found_sections.append(section_name)
                else:
                    validation_result["missing_sections"].append(section_name)
            
            validation_result["sections_found"] = found_sections
            
            # Determine validity
            if validation_result["missing_sections"]:
                validation_result["warnings"].append(
                    f"Missing sections: {', '.join(validation_result['missing_sections'])}"
                )
            
            # Check for basic formatting
            if len(doc.paragraphs) < 10:
                validation_result["warnings"].append("Template seems too short - check formatting")
            
            self.log_info(f"Template validation completed: {len(found_sections)} sections found")
            
        except Exception as e:
            validation_result["is_valid"] = False
            validation_result["errors"].append(f"Template validation failed: {str(e)}")
            self.log_error(f"Template validation error: {str(e)}")
        
        return validation_result
    
    def extract_template_info(self, template_path: Path) -> Dict[str, Any]:
        """Extract information about the template"""
        
        info = {
            "file_path": str(template_path),
            "file_size": template_path.stat().st_size if template_path.exists() else 0,
            "last_modified": datetime.fromtimestamp(template_path.stat().st_mtime) if template_path.exists() else None,
            "sections": [],
            "paragraphs_count": 0,
            "styles_used": []
        }
        
        try:
            # Load document
            doc = Document(str(template_path))
            
            # Count paragraphs
            info["paragraphs_count"] = len(doc.paragraphs)
            
            # Extract sections
            sections = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) < 50 and text.isupper():
                    sections.append(text)
            
            info["sections"] = sections
            
            # Extract styles
            styles = set()
            for paragraph in doc.paragraphs:
                if paragraph.style:
                    styles.add(paragraph.style.name)
            
            info["styles_used"] = list(styles)
            
        except Exception as e:
            self.log_error(f"Failed to extract template info: {str(e)}")
        
        return info

    def _save_without_overwrite(self, doc, output_file: Path):
        """CRITICAL SAFETY RULE: NEVER OVERWRITE EXISTING FILES

        This function implements the mandatory no-overwrite policy:
        - If file exists: Generate unique filename with timestamp
        - If file doesn't exist: Save with original filename
        - Always preserves existing files to prevent data loss
        """
        import datetime

        if output_file.exists():
            # File exists - generate unique filename to avoid overwrite
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            stem = output_file.stem
            suffix = output_file.suffix
            unique_filename = f"{stem}_{timestamp}{suffix}"
            unique_file = output_file.parent / unique_filename

            doc.save(str(unique_file))
            self.log_warning(f"⚠️ FILE EXISTS - Saved as: {unique_filename}")
            self.log_warning(f"   Original file preserved: {output_file.name}")
            self.log_info(f"   New file: {unique_file}")

            return unique_file
        else:
            # File doesn't exist - save normally
            doc.save(str(output_file))
            self.log_info(f"✅ File saved: {output_file.name}")
            return output_file
