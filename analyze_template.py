#!/usr/bin/env python3
"""
Script to analyze the DOCX template and show its structure
"""

import sys
from pathlib import Path
from docx import Document

def analyze_template(template_path):
    """Analyze the DOCX template structure"""
    
    print("=" * 80)
    print("📄 TEMPLATE ANALYSIS")
    print("=" * 80)
    
    try:
        doc = Document(template_path)
        
        print(f"📁 Template: {template_path}")
        print(f"📊 Total paragraphs: {len(doc.paragraphs)}")
        print(f"📊 Total sections: {len(doc.sections)}")
        print()
        
        print("📝 PARAGRAPH CONTENT:")
        print("-" * 80)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # Only show non-empty paragraphs
                style_name = paragraph.style.name if paragraph.style else "No Style"
                print(f"[{i:3d}] Style: {style_name:<20} | Text: {text[:100]}{'...' if len(text) > 100 else ''}")
        
        print()
        print("🔍 SEARCHING FOR KEY SECTIONS:")
        print("-" * 80)
        
        # Search for common CV section headers
        section_keywords = [
            'CURRICULUM', 'CV', 'RESUME', 'CURRÍCULUM',
            'PROFESSIONAL', 'SUMMARY', 'EXECUTIVE', 'RESUMEN',
            'SKILLS', 'COMPETENCIAS', 'TECHNICAL', 'SOFTWARE',
            'EXPERIENCE', 'EXPERIENCIA', 'EMPLOYMENT', 'WORK',
            'EDUCATION', 'EDUCACIÓN', 'FORMACIÓN',
            'PROJECTS', 'PROYECTOS'
        ]
        
        found_sections = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip().upper()
            for keyword in section_keywords:
                if keyword in text:
                    found_sections.append((i, paragraph.style.name, paragraph.text.strip()))
                    break
        
        if found_sections:
            print("✅ Found potential section headers:")
            for i, style, text in found_sections:
                print(f"   [{i:3d}] {style:<20} | {text}")
        else:
            print("❌ No clear section headers found")
        
        print()
        print("🎨 STYLES USED:")
        print("-" * 80)
        
        styles = set()
        for paragraph in doc.paragraphs:
            if paragraph.style:
                styles.add(paragraph.style.name)
        
        for style in sorted(styles):
            print(f"   • {style}")
        
        print()
        print("=" * 80)
        
    except Exception as e:
        print(f"❌ Error analyzing template: {str(e)}")

if __name__ == "__main__":
    template_path = Path("templates/PedroHerrera_PA_SaaS_B2B_Remote_2025.docx")
    
    if not template_path.exists():
        print(f"❌ Template not found: {template_path}")
        sys.exit(1)
    
    analyze_template(template_path)
