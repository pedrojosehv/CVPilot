"""
Script para diagnosticar qué está pasando con la detección de títulos en el template
"""

from docx import Document
import os

def diagnose_template():
    print("🔍 DIAGNÓSTICO DEL TEMPLATE ACTUAL")
    print("=" * 50)

    # Ruta del template actual
    template_path = "templates/PedroHerrera_PA_SaaS_B2B_Remote_2025.docx"

    if not os.path.exists(template_path):
        print(f"❌ Template no encontrado: {template_path}")
        return

    print(f"✅ Template encontrado: {template_path}")

    try:
        doc = Document(template_path)
        print(f"\n📊 Total de párrafos: {len(doc.paragraphs)}")

        print("\n📋 ANÁLISIS DE CONTENIDO:")
        print("-" * 30)

        potential_titles = []
        experience_section_found = False

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            if text:
                print(f"[{i:2d}] {text}")

                # Detectar sección de experiencia
                if any(keyword in text.upper() for keyword in ['EXPERIENCE', 'PROFESSIONAL']):
                    experience_section_found = True
                    print(f"      📍 SECCIÓN DE EXPERIENCIA ENCONTRADA")

                # Buscar patrones de títulos potenciales
                if experience_section_found:
                    # Patrones que indican títulos de experiencia
                    if any(pattern in text.lower() for pattern in [
                        'analyst', 'manager', 'specialist', 'consultant', 'developer',
                        'engineer', 'director', 'coordinator', 'lead', 'senior', 'junior'
                    ]):
                        # Verificar si tiene fechas
                        if any(char in text for char in ['/', '-', 'present', '2020', '2021', '2022', '2023']):
                            potential_titles.append((i, text))
                            print(f"      🎯 POSIBLE TÍTULO DETECTADO")

        print("
🎯 TÍTULOS POTENCIALES ENCONTRADOS:"        print("-" * 40)

        if potential_titles:
            for i, (line_num, title) in enumerate(potential_titles, 1):
                print(f"{i}. Línea {line_num}: '{title}'")

                # Analizar el título
                parts = title.split('\t') if '\t' in title else title.split('   ')
                if len(parts) > 1:
                    title_part = parts[0].strip()
                    date_part = parts[1].strip()
                    print(f"   📝 Título: '{title_part}'")
                    print(f"   📅 Fecha: '{date_part}'")
                else:
                    print("   ⚠️  No se pudo separar título de fecha")

                print()
        else:
            print("❌ NO SE ENCONTRARON TÍTULOS POTENCIALES")
            print("\n💡 Esto explica por qué el sistema dice 'No experience job titles found'")

        print("
📊 RESUMEN:"        print("-" * 20)
        print(f"✅ Template tiene {len(doc.paragraphs)} párrafos")
        print(f"✅ Sección de experiencia: {'SÍ' if experience_section_found else 'NO'}")
        print(f"✅ Títulos potenciales encontrados: {len(potential_titles)}")

        if not potential_titles:
            print("\n🎯 CONCLUSIÓN: El template actual no tiene títulos específicos de experiencia")
            print("💡 Para que el sistema funcione, necesitas títulos como:")
            print("   'Product Analyst 11/2023 - Present'")
            print("   'Project Manager 08/2022 - 11/2023'")
            print("   'Quality Assurance Analyst 11/2021 - 08/2022'")

    except Exception as e:
        print(f"❌ Error al leer el template: {e}")

if __name__ == "__main__":
    diagnose_template()
