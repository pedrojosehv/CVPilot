#!/usr/bin/env python3
"""
Script para verificar los cambios en los nombres de los roles en el CV generado
"""

from docx import Document
from pathlib import Path

def verify_role_changes():
    """Verificar los cambios en los nombres de los roles"""
    
    print("🔍 VERIFICANDO CAMBIOS EN NOMBRES DE ROLES")
    print("=" * 60)
    
    # Ruta del CV generado
    cv_path = Path("output/Project Manager - General - Excel (advanced), Amazon Seller Central, Helium 10/PedroHerrera_PJM_GEN_MP_2025.docx")
    
    if not cv_path.exists():
        print(f"❌ No se encontró el CV: {cv_path}")
        return
    
    try:
        doc = Document(str(cv_path))
        
        print(f"✅ CV cargado: {cv_path}")
        print(f"📊 Total de párrafos: {len(doc.paragraphs)}")
        
        # Buscar cambios específicos
        print(f"\n📋 ANÁLISIS DE CAMBIOS:")
        print("-" * 40)
        
        cv_title_found = False
        experience_titles_found = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Buscar el título principal del CV
            if "PROJECT MANAGER" in text.upper() and len(text) < 50:
                cv_title_found = True
                print(f"🎯 TÍTULO PRINCIPAL DEL CV (línea {i+1}):")
                print(f"   '{text}'")
                print()
            
            # Buscar títulos de experiencia
            if "Project Manager —" in text and ("GCA" in text or "Taime" in text or "Loszen" in text):
                experience_titles_found.append({
                    'line': i+1,
                    'text': text
                })
        
        # Mostrar títulos de experiencia encontrados
        if experience_titles_found:
            print(f"💼 TÍTULOS DE EXPERIENCIA ACTUALIZADOS:")
            for title in experience_titles_found:
                print(f"   Línea {title['line']}: '{title['text']}'")
        else:
            print(f"⚠️ No se encontraron títulos de experiencia con el formato esperado")
        
        # Verificar si se aplicaron los cambios
        print(f"\n📊 RESUMEN DE CAMBIOS:")
        print("-" * 40)
        
        if cv_title_found:
            print(f"✅ Título principal del CV: ACTUALIZADO a 'PROJECT MANAGER'")
        else:
            print(f"❌ Título principal del CV: NO ENCONTRADO")
        
        if experience_titles_found:
            print(f"✅ Títulos de experiencia: {len(experience_titles_found)} actualizados")
            for title in experience_titles_found:
                company = "Desconocida"
                if "GCA" in title['text']:
                    company = "GCA"
                elif "Taime" in title['text']:
                    company = "Industrias de Tapas Taime"
                elif "Loszen" in title['text']:
                    company = "Loszen"
                print(f"   - {company}: '{title['text']}'")
        else:
            print(f"❌ Títulos de experiencia: NO ACTUALIZADOS")
        
        # Comparar con el bullet_pool
        print(f"\n🔍 COMPARACIÓN CON BULLET_POOL:")
        print("-" * 40)
        
        bullet_pool_path = Path("templates/bullet_pool.docx")
        if bullet_pool_path.exists():
            try:
                bullet_doc = Document(str(bullet_pool_path))
                
                # Buscar títulos disponibles en bullet_pool
                available_titles = []
                current_profile = None
                current_company = None
                
                for paragraph in bullet_doc.paragraphs:
                    text = paragraph.text.strip()
                    
                    if "PERFIL AVANZADO" in text.upper():
                        current_profile = "AVANZADO"
                    elif "PERFIL BÁSICO" in text.upper():
                        current_profile = "BÁSICO"
                    elif "GCA" in text or "Growing Companies Advisors" in text:
                        current_company = "GCA"
                    elif "Industrias de Tapas Taime" in text or "Taime" in text:
                        current_company = "Industrias de Tapas Taime"
                    elif "Loszen" in text:
                        current_company = "Loszen"
                    
                    # Buscar títulos de roles
                    if current_profile and current_company:
                        if any(keyword in text for keyword in [
                            "Manager", "Director", "Analyst", "Specialist", "Coordinator",
                            "Lead", "Senior", "Junior", "Principal", "Head", "Chief"
                        ]):
                            if len(text) < 100 and not text.startswith('•'):
                                available_titles.append({
                                    'title': text,
                                    'company': current_company,
                                    'profile': current_profile
                                })
                
                # Mostrar títulos disponibles
                if available_titles:
                    print(f"📋 TÍTULOS DISPONIBLES EN BULLET_POOL:")
                    for title in available_titles:
                        print(f"   - {title['company']} ({title['profile']}): '{title['title']}'")
                    
                    # Verificar si el título usado está en la lista
                    used_title = "Project Manager"
                    found_in_pool = any(used_title.lower() in title['title'].lower() for title in available_titles)
                    
                    if found_in_pool:
                        print(f"\n✅ El título '{used_title}' SÍ está disponible en bullet_pool")
                    else:
                        print(f"\n❌ El título '{used_title}' NO está disponible en bullet_pool")
                        print(f"   Títulos disponibles: {[title['title'] for title in available_titles]}")
                else:
                    print(f"❌ No se encontraron títulos en bullet_pool")
                    
            except Exception as e:
                print(f"❌ Error leyendo bullet_pool: {e}")
        else:
            print(f"❌ No se encontró bullet_pool.docx")
        
        print(f"\n✅ Verificación completada")
        
    except Exception as e:
        print(f"❌ Error al verificar el CV: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    verify_role_changes()
