#!/usr/bin/env python3
"""
Debug Template Content - Analyze raw content of updated CV template
CVPilot - Debug why Business Operations titles are not being detected
"""

from pathlib import Path
from docx import Document
from rich.console import Console
from rich.panel import Panel

console = Console()

def debug_template_content():
    """Debug the raw content of the CV template to understand why Business Operations is not detected"""

    console.print("[bold blue]🔍 DEBUGGING TEMPLATE CONTENT[/bold blue]")
    console.print("[bold yellow]Analyzing raw content to find Business Operations titles[/bold yellow]")
    console.print("=" * 70)

    template_path = Path("templates/PedroHerrera_PA_SaaS_B2B_Remote_2025.docx")

    if not template_path.exists():
        console.print(f"[red]❌ Template not found: {template_path}[/red]")
        return

    console.print(f"\n[cyan]📄 Analyzing template: {template_path.name}[/cyan]")

    try:
        doc = Document(template_path)
        console.print(f"✅ Template loaded successfully ({len(doc.paragraphs)} paragraphs)")

        # Extract all text content
        all_text = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                all_text.append(f"[{i:2d}] {text}")
                # Check for Business Operations mentions
                if 'Business Operations' in text:
                    console.print(f"[green]🎯 FOUND 'Business Operations' in line {i}:[/green]")
                    console.print(f"   [bold]{text}[/bold]")

        # Show all text content
        console.print("\n[cyan]📝 RAW TEMPLATE CONTENT:[/cyan]")
        full_text = "\n".join(all_text)

        panel = Panel(
            "\n".join(all_text[:20]),  # Show first 20 lines
            title="First 20 Lines of Template",
            border_style="blue"
        )
        console.print(panel)

        if len(all_text) > 20:
            console.print(f"[yellow]... and {len(all_text) - 20} more lines[/yellow]")

        # Search for job titles
        console.print("\n[cyan]🔍 SEARCHING FOR JOB TITLES:[/cyan]")

        job_titles = [
            'Business Operations',
            'Product Analyst',
            'Digital Product Specialist',
            'Quality Assurance Analyst',
            'Project Manager',
            'Product Manager',
            'Business Analyst',
            'Operations Specialist'
        ]

        found_titles = []
        for title in job_titles:
            if title.lower() in full_text.lower():
                found_titles.append(title)
                console.print(f"[green]✅ Found: {title}[/green]")
            else:
                console.print(f"[red]❌ Not found: {title}[/red]")

        # Search for date patterns
        console.print("\n[cyan]🔍 SEARCHING FOR DATE PATTERNS:[/cyan]")
        import re

        date_patterns = [
            r'\d{1,2}/\d{4}\s*[-–]\s*\d{1,2}/\d{4}',
            r'\d{1,2}/\d{4}\s*[-–]\s*Present',
            r'\d{4}\s*[-–]\s*\d{4}',
            r'\d{4}\s*[-–]\s*Present'
        ]

        for pattern in date_patterns:
            matches = re.findall(pattern, full_text)
            if matches:
                console.print(f"[green]✅ Found date pattern '{pattern}': {matches}[/green]")
            else:
                console.print(f"[red]❌ No matches for pattern '{pattern}'[/red]")

        # Summary
        console.print("\n[bold green]📊 DEBUG SUMMARY:[/bold green]")
        console.print(f"   • Total paragraphs: {len(doc.paragraphs)}")
        console.print(f"   • Lines with text: {len(all_text)}")
        console.print(f"   • Business Operations found: {'✅ YES' if 'Business Operations' in full_text else '❌ NO'}")
        console.print(f"   • Job titles found: {len(found_titles)}")
        console.print(f"   • Date patterns found: {'✅ YES' if any(re.findall('|'.join(date_patterns), full_text)) else '❌ NO'}")

        if 'Business Operations' not in full_text:
            console.print("\n[yellow]💡 The template may not contain 'Business Operations' titles yet,[/yellow]")
            console.print("[yellow]   or they may be formatted differently than expected.[/yellow]")

    except Exception as e:
        console.print(f"[red]❌ Error analyzing template: {e}[/red]")

if __name__ == "__main__":
    debug_template_content()
