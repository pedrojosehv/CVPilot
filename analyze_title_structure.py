#!/usr/bin/env python3
"""
Script para analizar la estructura de títulos de roles en bullet_pool.docx
"""

from docx import Document
from pathlib import Path

def analyze_title_structure():
    """Analizar la estructura de títulos de roles en bullet_pool.docx"""
    
    print("🔍 ANÁLISIS DE ESTRUCTURA DE TÍTULOS DE ROLES")
    print("=" * 60)
    
    bullet_pool_path = Path("templates/bullet_pool.docx")
    
    if not bullet_pool_path.exists():
        print(f"❌ No se encontró el archivo: {bullet_pool_path}")
        return
    
    try:
        doc = Document(str(bullet_pool_path))
        
        print(f"✅ Archivo cargado: {bullet_pool_path}")
        print(f"📊 Total de párrafos: {len(doc.paragraphs)}")
        print(f"📊 Total de tablas: {len(doc.tables)}")
        
        # Analizar párrafos para encontrar títulos de roles
        print(f"\n📋 ANÁLISIS DE PÁRRAFOS:")
        print("-" * 40)
        
        titles_found = []
        current_company = None
        current_profile = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                continue
                
            # Buscar perfiles
            if "PERFIL AVANZADO" in text.upper():
                current_profile = "AVANZADO"
                print(f"📍 Perfil Avanzado encontrado en línea {i+1}")
                continue
                
            if "PERFIL BÁSICO" in text.upper():
                current_profile = "BÁSICO"
                print(f"📍 Perfil Básico encontrado en línea {i+1}")
                continue
            
            # Buscar nombres de empresas
            company_indicators = [
                "GCA", "Growing Companies Advisors",
                "Industrias de Tapas Taime", "Taime",
                "Loszen", "Loszen Tech",
                "Otros", "Otras empresas"
            ]
            
            for indicator in company_indicators:
                if indicator.lower() in text.lower():
                    current_company = indicator
                    print(f"🏢 Empresa encontrada: {indicator} (línea {i+1})")
                    break
            
            # Buscar títulos de roles (patrones específicos)
            if any(keyword in text for keyword in [
                "Manager", "Director", "Analyst", "Specialist", "Coordinator",
                "Lead", "Senior", "Junior", "Principal", "Head", "Chief"
            ]):
                if len(text) < 100 and not text.startswith('•'):
                    titles_found.append({
                        'text': text,
                        'line': i+1,
                        'company': current_company,
                        'profile': current_profile
                    })
                    print(f"💼 Título encontrado: {text} (línea {i+1}, {current_company}, {current_profile})")
        
        # Analizar tablas para encontrar títulos de roles
        print(f"\n📊 ANÁLISIS DE TABLAS:")
        print("-" * 40)
        
        for table_idx, table in enumerate(doc.tables):
            print(f"📋 Tabla {table_idx + 1}: {len(table.rows)} filas, {len(table.columns)} columnas")
            
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                
                if row_text:
                    print(f"   Fila {row_idx + 1}: {' | '.join(row_text)}")
                    
                    # Buscar títulos de roles en las celdas
                    for cell_text in row_text:
                        if any(keyword in cell_text for keyword in [
                            "Manager", "Director", "Analyst", "Specialist", "Coordinator",
                            "Lead", "Senior", "Junior", "Principal", "Head", "Chief"
                        ]):
                            if len(cell_text) < 100:
                                titles_found.append({
                                    'text': cell_text,
                                    'line': f"Tabla{table_idx+1}-Fila{row_idx+1}",
                                    'company': "GCA (Tabla)",
                                    'profile': "AVANZADO (Tabla)"
                                })
                                print(f"   💼 Título en tabla: {cell_text}")
        
        # Resumen de títulos encontrados
        print(f"\n📈 RESUMEN DE TÍTULOS ENCONTRADOS:")
        print("-" * 40)
        
        if titles_found:
            # Agrupar por empresa
            companies = {}
            for title in titles_found:
                company = title['company'] or "Sin empresa"
                if company not in companies:
                    companies[company] = []
                companies[company].append(title)
            
            for company, titles in companies.items():
                print(f"\n🏢 {company}:")
                for title in titles:
                    profile = title['profile'] or "Sin perfil"
                    print(f"   • {title['text']} ({profile})")
        else:
            print("❌ No se encontraron títulos de roles")
        
        print(f"\n✅ Análisis completado")
        
    except Exception as e:
        print(f"❌ Error al analizar el archivo: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_title_structure()
