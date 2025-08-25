#!/usr/bin/env python3
"""
Script para leer y analizar el contenido del bullet_pool.docx
"""

from docx import Document
from pathlib import Path

def read_bullet_pool():
    print('🔍 ANALIZANDO BULLET_POOL.DOCX')
    print('=' * 60)
    
    bullet_pool_path = Path('templates/bullet_pool.docx')
    
    if not bullet_pool_path.exists():
        print(f'❌ No se encontró: {bullet_pool_path}')
        return
    
    doc = Document(str(bullet_pool_path))
    
    print(f'📄 Archivo: {bullet_pool_path.name}')
    print(f'📊 Total de párrafos: {len(doc.paragraphs)}')
    print()
    
    print('📋 CONTENIDO COMPLETO:')
    print('=' * 60)
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f'📝 Línea {i+1}: {text}')
    
    print()
    print('🎯 ANÁLISIS DE TABLAS:')
    print('=' * 60)
    
    for i, table in enumerate(doc.tables):
        print(f'📊 Tabla {i+1}:')
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                print(f'   Fila {row_idx+1}: {" | ".join(row_text)}')
        print()

if __name__ == "__main__":
    read_bullet_pool()
