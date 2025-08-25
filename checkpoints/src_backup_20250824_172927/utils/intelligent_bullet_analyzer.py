"""
Enhanced Intelligent Bullet Analyzer for CVPilot
Uses LLM to analyze job descriptions and select optimal bullets from pool
Supports advanced profile detection with role progression and company context
"""

import re
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
import docx

from .api_key_manager import get_api_key, mark_api_error
from .logger import LoggerMixin

class EnhancedBulletAnalyzer(LoggerMixin):
    """Enhanced analyzer with support for advanced profiles and role progression"""
    
    def __init__(self, bullet_pool_path: str = "templates/bullet_pool.docx"):
        super().__init__()
        self.bullet_pool_path = bullet_pool_path
        self.bullet_pool = self._load_enhanced_bullet_pool()
        self.role_progression = self._extract_role_progression()
        self.company_contexts = self._extract_company_contexts()
        self.role_titles = self._extract_role_titles()
        
    def _load_enhanced_bullet_pool(self) -> Dict[str, Any]:
        """Load enhanced bullet pool with profiles, roles, and tables"""
        try:
            doc = docx.Document(self.bullet_pool_path)
            
            # Extract structured data
            profiles = {
                "advanced": {
                    "gca_roles": self._extract_gca_roles_from_tables(doc),
                    "bullets": self._extract_bullets_from_paragraphs(doc)
                },
                "basic": {
                    "companies": self._extract_other_companies(doc),
                    "bullets": self._extract_bullets_from_paragraphs(doc)
                }
            }
            
            return profiles
            
        except Exception as e:
            self.log_error(f"Error loading enhanced bullet pool: {e}")
            return {"advanced": {"gca_roles": [], "bullets": {}}, "basic": {"companies": [], "bullets": {}}}
    
    def _extract_gca_roles_from_tables(self, doc) -> List[Dict[str, Any]]:
        """Extract GCA role progression from tables"""
        roles = []
        
        for table in doc.tables:
            if table.rows:
                row = table.rows[0]
                if len(row.cells) >= 2:
                    role_text = row.cells[0].text.strip()
                    period = row.cells[1].text.strip()
                    
                    # Parse multiple roles if separated by |
                    role_titles = [r.strip() for r in role_text.split('|')]
                    
                    # Extract application context if present
                    app_context = ""
                    if '(' in role_text and ')' in role_text:
                        app_context = role_text[role_text.find('(')+1:role_text.find(')')]
                    
                    roles.append({
                        "role_titles": role_titles,
                        "period": period,
                        "application_context": app_context,
                        "primary_role": role_titles[0] if role_titles else "",
                        "is_current": "Present" in period
                    })
        
        # Sort by period (most recent first)
        roles.sort(key=lambda x: x["is_current"], reverse=True)
        return roles
    
    def _extract_bullets_from_paragraphs(self, doc) -> Dict[str, List[str]]:
        """Extract bullets organized by company"""
        bullets = {}
        current_company = None
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Detect company headers
            if "—" in text and ("Remote" in text or "," in text):
                current_company = text
                bullets[current_company] = []
            elif current_company and not any(keyword in text for keyword in ["—", "Perfil"]):
                # This is a bullet point
                bullets[current_company].append(text)
        
        return bullets
    
    def _extract_other_companies(self, doc) -> List[Dict[str, str]]:
        """Extract other companies (non-GCA)"""
        companies = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "—" in text and ("Remote" in text or "," in text) and "GCA" not in text:
                # Parse company info
                parts = text.split("—")
                if len(parts) >= 2:
                    company_part = parts[0].strip()
                    location_part = parts[1].strip()
                    
                    # Extract company name and industry
                    if "," in company_part:
                        company_info = company_part.split(",")
                        name = company_info[0].strip()
                        industry = company_info[1].strip() if len(company_info) > 1 else ""
                    else:
                        name = company_part
                        industry = ""
                    
                    companies.append({
                        "name": name,
                        "industry": industry,
                        "location": location_part,
                        "full_text": text
                    })
        
        return companies
    
    def _extract_role_progression(self) -> List[Dict[str, Any]]:
        """Extract role progression from advanced profile"""
        if "advanced" in self.bullet_pool and "gca_roles" in self.bullet_pool["advanced"]:
            return self.bullet_pool["advanced"]["gca_roles"]
        return []
    
    def _extract_company_contexts(self) -> Dict[str, Dict[str, str]]:
        """Extract company contexts for better matching"""
        contexts = {}
        
        # GCA context
        contexts["GCA"] = {
            "industry": "Consulting",
            "type": "Consulting firm",
            "location": "U.S. (Remote)",
            "specialization": "SaaS B2B Fintech, Product Management",
            "size": "Small-Medium"
        }
        
        # Other companies from basic profile
        if "basic" in self.bullet_pool and "companies" in self.bullet_pool["basic"]:
            for company in self.bullet_pool["basic"]["companies"]:
                contexts[company["name"]] = {
                    "industry": company["industry"],
                    "location": company["location"],
                    "type": company["industry"],
                    "specialization": self._infer_specialization(company["industry"]),
                    "size": self._infer_size(company["industry"])
                }
        
        return contexts
    
    def _extract_role_titles(self) -> Dict[str, Dict[str, List[str]]]:
        """Extract role titles organized by profile and company"""
        try:
            doc = docx.Document(self.bullet_pool_path)
            
            titles = {
                "advanced": {},
                "basic": {}
            }
            
            current_profile = None
            current_company = None
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                # Detect profile
                if "PERFIL AVANZADO" in text.upper():
                    current_profile = "advanced"
                    continue
                    
                if "PERFIL BÁSICO" in text.upper():
                    current_profile = "basic"
                    continue
                
                # Detect company
                if "GCA" in text or "Growing Companies Advisors" in text:
                    current_company = "GCA"
                elif "Industrias de Tapas Taime" in text or "Taime" in text:
                    current_company = "Industrias de Tapas Taime"
                elif "Loszen" in text:
                    current_company = "Loszen"
                
                # Extract role titles from paragraphs
                if current_profile and current_company:
                    if any(keyword in text for keyword in [
                        "Manager", "Director", "Analyst", "Specialist", "Coordinator",
                        "Lead", "Senior", "Junior", "Principal", "Head", "Chief"
                    ]):
                        if len(text) < 100 and not text.startswith('•'):
                            if current_company not in titles[current_profile]:
                                titles[current_profile][current_company] = []
                            titles[current_profile][current_company].append(text)
            
            # Extract from tables (GCA advanced roles) - CORREGIDO
            for table_idx, table in enumerate(doc.tables):
                if table.rows:
                    row = table.rows[0]
                    if len(row.cells) >= 2:
                        role_text = row.cells[0].text.strip()
                        
                        # Parse multiple roles if separated by |
                        role_titles = [r.strip() for r in role_text.split('|')]
                        
                        # Clean up role titles
                        clean_titles = []
                        for title in role_titles:
                            # Remove application context in parentheses
                            if '(' in title and ')' in title:
                                title = title[:title.find('(')].strip()
                            # Remove period information if present
                            if '\t' in title:
                                title = title.split('\t')[0].strip()
                            clean_titles.append(title)
                        
                        # Add to GCA advanced titles
                        if "GCA" not in titles["advanced"]:
                            titles["advanced"]["GCA"] = []
                        titles["advanced"]["GCA"].extend(clean_titles)
                        
                        self.log_info(f"📋 Extracted from table {table_idx + 1}: {clean_titles}")
            
            # Log extracted titles for debugging
            self.log_info(f"📊 Extracted role titles:")
            for profile, companies in titles.items():
                for company, role_list in companies.items():
                    self.log_info(f"   {profile} - {company}: {role_list}")
            
            return titles
            
        except Exception as e:
            self.log_error(f"Error extracting role titles: {e}")
            return {"advanced": {}, "basic": {}}
    
    def select_appropriate_role_title(self, job_data: Dict[str, Any], profile_type: str) -> str:
        """Select appropriate role title based on job data and profile type"""
        
        job_title = job_data.get('job_title', '').lower()
        company = job_data.get('company', '').lower()
        
        # Determine which company to use for title selection
        target_company = None
        
        # Map job companies to bullet pool companies
        if 'gca' in company or 'growing companies' in company:
            target_company = "GCA"
        elif 'taime' in company or 'tapas' in company or 'industrias' in company:
            target_company = "Industrias de Tapas Taime"
        elif 'loszen' in company:
            target_company = "Loszen"
        else:
            # Default to GCA for unknown companies
            target_company = "GCA"
        
        # Get available titles for the profile and company
        available_titles = []
        if profile_type in self.role_titles and target_company in self.role_titles[profile_type]:
            available_titles = self.role_titles[profile_type][target_company]
        
        if not available_titles:
            # Fallback to GCA titles
            if profile_type in self.role_titles and "GCA" in self.role_titles[profile_type]:
                available_titles = self.role_titles[profile_type]["GCA"]
        
        if not available_titles:
            # Final fallback - generate based on job title
            return self._generate_fallback_title(job_title)
        
        # Select the most appropriate title based on job requirements
        selected_title = self._select_best_matching_title(job_title, available_titles)
        
        self.log_info(f"🎯 Selected role title: {selected_title} (from {target_company}, {profile_type} profile)")
        return selected_title
    
    def _select_best_matching_title(self, job_title: str, available_titles: List[str]) -> str:
        """Select the best matching title from available options"""
        
        # Define priority keywords for different job types
        priority_keywords = {
            'product': ['Product Manager', 'Product Owner', 'Product Analyst'],
            'project': ['Project Manager', 'Project Coordinator'],
            'business': ['Business Analyst', 'Business Specialist'],
            'data': ['Data Analyst', 'Quality Analyst', 'Quality Assurance'],
            'operations': ['Operations Specialist', 'Operations Manager'],
            'growth': ['Growth Manager', 'Product Manager'],
            'technical': ['Technical Manager', 'Product Manager'],
            'marketplace': ['Project Manager', 'Product Manager']
        }
        
        # Find the best match
        for category, preferred_titles in priority_keywords.items():
            if category in job_title:
                for preferred in preferred_titles:
                    for available in available_titles:
                        if preferred.lower() in available.lower():
                            return available
        
        # If no specific match, return the first available title
        return available_titles[0] if available_titles else "Professional"
    
    def _generate_fallback_title(self, job_title: str) -> str:
        """Generate a fallback title based on job title"""
        
        if 'manager' in job_title:
            return "Professional Manager"
        elif 'analyst' in job_title:
            return "Professional Analyst"
        elif 'specialist' in job_title:
            return "Professional Specialist"
        elif 'coordinator' in job_title:
            return "Professional Coordinator"
        else:
            return "Professional"
    
    def _infer_specialization(self, industry: str) -> str:
        """Infer specialization based on industry"""
        industry_lower = industry.lower()
        if "manufacturing" in industry_lower:
            return "Manufacturing Operations, Quality Control"
        elif "startup" in industry_lower or "app" in industry_lower:
            return "Mobile Development, User Experience"
        elif "consulting" in industry_lower:
            return "Business Strategy, Process Optimization"
        else:
            return "General Business Operations"
    
    def _infer_size(self, industry: str) -> str:
        """Infer company size based on industry"""
        industry_lower = industry.lower()
        if "startup" in industry_lower:
            return "Small"
        elif "manufacturing" in industry_lower:
            return "Medium-Large"
        else:
            return "Medium"
    
    def _parse_bullet_pool(self, content: List[str]) -> Dict[str, List[str]]:
        """Parse bullet pool content into structured format"""
        bullet_pool = {}
        current_company = None
        current_bullets = []
        
        for line in content:
            line = line.strip()
            if not line:
                continue
                
            # Check if this is a company header (contains "—")
            if "—" in line and ("Remote" in line or "," in line):
                # Save previous company bullets
                if current_company and current_bullets:
                    bullet_pool[current_company] = current_bullets
                
                # Start new company
                current_company = line
                current_bullets = []
                
            # Check if this is a bullet point (starts with number or contains key metrics)
            elif (line[0].isdigit() or 
                  any(keyword in line.lower() for keyword in ['led', 'drove', 'achieved', 'increased', 'reduced', 'spearheaded', 'mitigated', 'resolved'])):
                current_bullets.append(line)
        
        # Save last company
        if current_company and current_bullets:
            bullet_pool[current_company] = current_bullets
            
        return bullet_pool
    
    def analyze_job_and_select_bullets(self, job_data: Dict[str, Any], 
                                     job_description: str = "") -> List[str]:
        """Enhanced analysis with role progression and company context"""
        
        if not self.bullet_pool:
            self.log_warning("No enhanced bullet pool available")
            return []
        
        # Determine optimal profile and role matching
        profile_type, matching_role = self._determine_optimal_profile(job_data)
        
        self.log_info(f"🎯 Selected profile: {profile_type}")
        if matching_role:
            self.log_info(f"🎯 Matching role: {matching_role.get('primary_role', 'Unknown')}")
        
        # Create enhanced prompt with role progression context
        prompt = self._create_enhanced_analysis_prompt(job_data, job_description, profile_type, matching_role)
        
        try:
            # Call LLM for intelligent analysis
            response = self._call_llm(prompt)
            
            # Parse LLM response to get selected bullets
            selected_bullets = self._parse_enhanced_llm_response(response, job_data, profile_type)
            
            self.log_info(f"✅ Selected {len(selected_bullets)} bullets using enhanced analysis")
            return selected_bullets
            
        except Exception as e:
            self.log_error(f"Enhanced LLM analysis failed: {e}")
            # Fallback to enhanced rule-based selection
            return self._enhanced_fallback_selection(job_data, profile_type)
    
    def _determine_optimal_profile(self, job_data: Dict[str, Any]) -> Tuple[str, Optional[Dict[str, Any]]]:
        """Determine whether to use advanced or basic profile"""
        
        job_title = job_data.get('job_title_original', '').lower()
        company = job_data.get('company', '').lower()
        skills = [s.lower() for s in job_data.get('skills', [])]
        
        # Check for advanced profile indicators
        advanced_indicators = [
            'product manager', 'product owner', 'product analyst', 'business analyst',
            'project manager', 'operations specialist', 'fintech', 'saas', 'ai', 'ml'
        ]
        
        advanced_score = sum(1 for indicator in advanced_indicators 
                           if any(indicator in text for text in [job_title, company] + skills))
        
        # If high advanced score, use advanced profile with role matching
        if advanced_score >= 2:
            matching_role = self._find_matching_gca_role(job_data)
            return "advanced", matching_role
        else:
            return "basic", None
    
    def _find_matching_gca_role(self, job_data: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Find the most matching GCA role based on job requirements"""
        
        job_title = job_data.get('job_title_original', '').lower()
        skills = [s.lower() for s in job_data.get('skills', [])]
        
        best_match = None
        best_score = 0
        
        for role in self.role_progression:
            role_score = 0
            
            # Score based on role titles
            for role_title in role.get('role_titles', []):
                role_title_lower = role_title.lower()
                if any(keyword in job_title for keyword in role_title_lower.split()):
                    role_score += 2
            
            # Score based on application context
            app_context = role.get('application_context', '').lower()
            context_keywords = ['saas', 'fintech', 'mobile', 'b2b', 'b2c', 'platform', 'operations']
            for keyword in context_keywords:
                if keyword in app_context and keyword in job_title:
                    role_score += 1
            
            # Prefer current role for recent experience
            if role.get('is_current', False):
                role_score += 1
            
            if role_score > best_score:
                best_score = role_score
                best_match = role
        
        return best_match
    
    def _create_enhanced_analysis_prompt(self, job_data: Dict[str, Any], job_description: str, 
                                       profile_type: str, matching_role: Optional[Dict[str, Any]]) -> str:
        """Create enhanced prompt with role progression context"""
        
        # Get available bullets based on profile type
        if profile_type == "advanced":
            bullets_source = self.bullet_pool.get("advanced", {}).get("bullets", {})
        else:
            bullets_source = self.bullet_pool.get("basic", {}).get("bullets", {})
        
        # Format bullet pool for prompt
        bullet_info = []
        for company, bullets in bullets_source.items():
            bullet_info.append(f"\n🏢 {company}:")
            for i, bullet in enumerate(bullets, 1):
                bullet_info.append(f"  {i}. {bullet}")
        
        # Role progression context
        role_context = ""
        if matching_role:
            role_context = f"""
MATCHED ROLE CONTEXT:
- Primary Role: {matching_role.get('primary_role', 'Unknown')}
- All Roles: {' | '.join(matching_role.get('role_titles', []))}
- Period: {matching_role.get('period', 'Unknown')}
- Application Context: {matching_role.get('application_context', 'Unknown')}
- Current Role: {'Yes' if matching_role.get('is_current') else 'No'}
"""
        
        prompt = f"""
You are an expert CV optimizer analyzing a job opportunity to select the most relevant achievement bullets.

TARGET JOB ANALYSIS:
- Job Title: {job_data.get('job_title_original', 'Unknown')}
- Company: {job_data.get('company', 'Unknown')}
- Required Skills: {', '.join(job_data.get('skills', [])[:10])}
- Required Software: {', '.join(job_data.get('software', [])[:8])}
- Seniority: {job_data.get('seniority', 'Unknown')}
- Industry Context: {self._infer_industry_from_job(job_data)}

{role_context}

PROFILE TYPE: {profile_type.upper()}

AVAILABLE BULLETS:
{''.join(bullet_info)}

SELECTION STRATEGY:
1. Prioritize bullets that align with the target role and industry
2. For ADVANCED profile: Focus on progression and leadership achievements
3. For BASIC profile: Focus on foundational skills and diverse experience
4. Include quantifiable results and specific technologies mentioned in job requirements
5. Maintain distribution: GCA (3 bullets), other companies (1-2 bullets each)

RESPONSE FORMAT:
Return ONLY a JSON object with this structure:
{{
    "selected_bullets": [
        "bullet text 1",
        "bullet text 2",
        ...
    ],
    "reasoning": "Brief explanation focusing on role alignment and progression",
    "profile_used": "{profile_type}"
}}

Select 5-7 most relevant bullets for this {job_data.get('job_title_original', 'role')} position:
"""
        
        return prompt
    
    def _infer_industry_from_job(self, job_data: Dict[str, Any]) -> str:
        """Infer industry context from job data"""
        company = job_data.get('company', '').lower()
        title = job_data.get('job_title_original', '').lower()
        skills = ' '.join(job_data.get('skills', [])).lower()
        
        # Industry keywords mapping
        industry_keywords = {
            "fintech": ["fintech", "financial", "banking", "payment", "crypto", "blockchain"],
            "saas": ["saas", "software", "platform", "cloud", "api"],
            "consulting": ["consulting", "advisory", "strategy", "transformation"],
            "manufacturing": ["manufacturing", "production", "industrial", "supply chain"],
            "startup": ["startup", "scale-up", "growth", "venture"],
            "enterprise": ["enterprise", "corporation", "large-scale", "global"]
        }
        
        text_to_analyze = f"{company} {title} {skills}"
        
        for industry, keywords in industry_keywords.items():
            if any(keyword in text_to_analyze for keyword in keywords):
                return industry.title()
        
        return "Technology"
    
    def _parse_enhanced_llm_response(self, response: str, job_data: Dict[str, Any], profile_type: str) -> List[str]:
        """Parse enhanced LLM response with better error handling"""
        try:
            import json
            
            # Clean response
            response = response.strip()
            if response.startswith('```json'):
                response = response[7:]
            if response.endswith('```'):
                response = response[:-3]
            
            # Parse JSON
            parsed = json.loads(response)
            
            selected_bullets = parsed.get("selected_bullets", [])
            reasoning = parsed.get("reasoning", "")
            
            self.log_info(f"🧠 LLM Reasoning: {reasoning}")
            self.log_info(f"📊 Profile used: {profile_type}")
            
            return selected_bullets[:7]  # Limit to 7 bullets max
            
        except Exception as e:
            self.log_error(f"Failed to parse enhanced LLM response: {e}")
            self.log_error(f"Raw response: {response[:200]}...")
            return []
    
    def _enhanced_fallback_selection(self, job_data: Dict[str, Any], profile_type: str) -> List[str]:
        """Enhanced rule-based fallback selection"""
        
        self.log_info(f"🔄 Using enhanced fallback selection for {profile_type} profile")
        
        selected_bullets = []
        job_title = job_data.get('job_title_original', '').lower()
        skills = [s.lower() for s in job_data.get('skills', [])]
        
        if profile_type == "advanced":
            bullets_source = self.bullet_pool.get("advanced", {}).get("bullets", {})
            # Prioritize GCA bullets for advanced profile
            gca_key = next((k for k in bullets_source.keys() if "GCA" in k), None)
            if gca_key and gca_key in bullets_source:
                gca_bullets = bullets_source[gca_key]
                # Select top 3 GCA bullets based on keyword matching
                scored_bullets = []
                for bullet in gca_bullets:
                    score = self._score_bullet_relevance(bullet, job_title, skills)
                    scored_bullets.append((bullet, score))
                
                scored_bullets.sort(key=lambda x: x[1], reverse=True)
                selected_bullets.extend([b[0] for b in scored_bullets[:3]])
        else:
            bullets_source = self.bullet_pool.get("basic", {}).get("bullets", {})
        
        # Add bullets from other companies
        for company, bullets in bullets_source.items():
            if "GCA" not in company and bullets:  # Skip GCA for basic or if already added
                # Select 1-2 bullets from each other company
                scored_bullets = []
                for bullet in bullets:
                    score = self._score_bullet_relevance(bullet, job_title, skills)
                    scored_bullets.append((bullet, score))
                
                scored_bullets.sort(key=lambda x: x[1], reverse=True)
                max_bullets = 2 if len(bullets) > 2 else len(bullets)
                selected_bullets.extend([b[0] for b in scored_bullets[:max_bullets]])
        
        return selected_bullets[:7]  # Limit to 7 total
    
    def _score_bullet_relevance(self, bullet: str, job_title: str, skills: List[str]) -> float:
        """Score bullet relevance based on job requirements"""
        bullet_lower = bullet.lower()
        score = 0.0
        
        # Score based on job title keywords
        job_keywords = job_title.split()
        for keyword in job_keywords:
            if keyword in bullet_lower and len(keyword) > 3:
                score += 2.0
        
        # Score based on skills
        for skill in skills:
            if skill in bullet_lower and len(skill) > 2:
                score += 1.5
        
        # Score based on quantifiable results
        if any(char in bullet for char in ['%', '$', 'MM', 'K']):
            score += 1.0
        
        # Score based on action verbs
        action_verbs = ['led', 'drove', 'achieved', 'increased', 'reduced', 'improved', 'developed', 'implemented']
        if any(verb in bullet_lower for verb in action_verbs):
            score += 0.5
        
        return score
    
    def _create_analysis_prompt(self, job_data: Dict[str, Any], job_description: str) -> str:
        """Create intelligent prompt for bullet analysis"""
        
        # Extract company information from bullet pool
        company_info = []
        for company, bullets in self.bullet_pool.items():
            company_info.append(f"COMPANY: {company}")
            company_info.append("BULLETS:")
            for i, bullet in enumerate(bullets, 1):
                company_info.append(f"{i}. {bullet}")
            company_info.append("")
        
        prompt = f"""
        You are an expert CV analyst. Analyze this job opportunity and select the most relevant bullets from the provided pool.

        JOB OPPORTUNITY:
        - Title: {job_data.get('job_title_original', 'Unknown')}
        - Company: {job_data.get('company', 'Unknown')}
        - Required Skills: {', '.join(job_data.get('skills', [])[:10])}
        - Required Software: {', '.join(job_data.get('software', [])[:5])}
        - Seniority: {job_data.get('seniority', 'Unknown')}
        
        JOB DESCRIPTION:
        {job_description if job_description else 'Not available'}

        BULLET POOL:
        {chr(10).join(company_info)}

        ANALYSIS REQUIREMENTS:
        1. For each company, analyze the role format: "Role 1 | Role 2 | ... | Role N (Specialization)"
        2. Select ONLY ONE role that best aligns with the target job
        3. Follow this distribution:
           - GCA (Growing Companies Advisors): 3 bullets
           - GCA (Second role): 2 bullets  
           - Industrias de Tapas Taime: 1 bullet
           - Loszen: 2 bullets
           - Industrias QProductos: 1 bullet
        4. Choose bullets that demonstrate:
           - Relevant technical skills
           - Quantifiable achievements
           - Industry-specific experience
           - Leadership and collaboration
           - Problem-solving abilities

        RESPONSE FORMAT:
        Return ONLY a JSON object with this structure:
        {{
            "selected_bullets": [
                "bullet text 1",
                "bullet text 2",
                ...
            ],
            "reasoning": "Brief explanation of why these bullets were selected"
        }}

        Select the most relevant bullets for this specific job opportunity:
        """
        
        return prompt
    
    def _call_llm(self, prompt: str) -> str:
        """Call LLM for intelligent analysis"""
        import openai
        
        max_retries = 3
        for attempt in range(max_retries):
            try:
                # Get API key from manager
                api_key = get_api_key("round_robin")
                if not api_key:
                    raise ValueError("No API key available")
                
                client = openai.OpenAI(api_key=api_key)
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": "You are an expert CV analyst and career consultant. Provide precise, relevant analysis for job applications."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=1000
                )
                return response.choices[0].message.content
                
            except Exception as e:
                error_msg = str(e).lower()
                
                # Check for rate limit errors
                if any(keyword in error_msg for keyword in ['rate limit', 'quota', 'too many requests']):
                    self.log_warning(f"Rate limit hit on attempt {attempt + 1}, rotating API key...")
                    mark_api_error(api_key, "rate_limit")
                    continue
                
                # Check for API key errors
                elif any(keyword in error_msg for keyword in ['invalid api key', 'authentication', 'unauthorized']):
                    self.log_error(f"API key error: {e}")
                    mark_api_error(api_key, "invalid_key")
                    continue
                
                # Other errors
                else:
                    self.log_error(f"LLM call failed: {e}")
                    if attempt == max_retries - 1:
                        break
                    continue
        
        # If all retries failed, return empty response
        return '{"selected_bullets": [], "reasoning": "LLM analysis failed"}'
    
    def _parse_llm_response(self, response: str, job_data: Dict[str, Any]) -> List[str]:
        """Parse LLM response to extract selected bullets"""
        import json
        
        try:
            # Try to parse JSON response
            parsed = json.loads(response)
            selected_bullets = parsed.get('selected_bullets', [])
            
            # Validate bullets exist in pool
            valid_bullets = []
            all_pool_bullets = []
            for bullets in self.bullet_pool.values():
                all_pool_bullets.extend(bullets)
            
            for bullet in selected_bullets:
                # Find matching bullet in pool (fuzzy matching)
                for pool_bullet in all_pool_bullets:
                    if self._bullet_matches(bullet, pool_bullet):
                        valid_bullets.append(pool_bullet)
                        break
            
            return valid_bullets
            
        except json.JSONDecodeError:
            self.log_error("Failed to parse LLM response as JSON")
            return self._fallback_bullet_selection(job_data)
    
    def _bullet_matches(self, selected: str, pool_bullet: str) -> bool:
        """Check if selected bullet matches pool bullet (fuzzy matching)"""
        # Remove numbers and clean up
        selected_clean = re.sub(r'^\d+\.\s*', '', selected.strip())
        pool_clean = re.sub(r'^\d+\.\s*', '', pool_bullet.strip())
        
        # Check for significant overlap
        selected_words = set(selected_clean.lower().split())
        pool_words = set(pool_clean.lower().split())
        
        # Calculate similarity
        intersection = selected_words.intersection(pool_words)
        union = selected_words.union(pool_words)
        
        if union:
            similarity = len(intersection) / len(union)
            return similarity > 0.6  # 60% similarity threshold
        
        return False
    
    def _fallback_bullet_selection(self, job_data: Dict[str, Any]) -> List[str]:
        """Fallback bullet selection using rule-based approach"""
        selected_bullets = []
        
        # Simple rule-based selection
        for company, bullets in self.bullet_pool.items():
            if "GCA" in company:
                # Select first 3 bullets for GCA
                selected_bullets.extend(bullets[:3])
            elif "Loszen" in company:
                # Select first 2 bullets for Loszen
                selected_bullets.extend(bullets[:2])
            elif "Taime" in company:
                # Select first 1 bullet for Taime
                selected_bullets.extend(bullets[:1])
            elif "QProductos" in company:
                # Select first 1 bullet for QProductos
                selected_bullets.extend(bullets[:1])
        
        self.log_info(f"✅ Fallback selection: {len(selected_bullets)} bullets")
        return selected_bullets
    
    def get_bullet_distribution(self) -> Dict[str, int]:
        """Get recommended bullet distribution"""
        return {
            "GCA": 3,
            "GCA_Second": 2,
            "Taime": 1,
            "Loszen": 2,
            "QProductos": 1
        }
