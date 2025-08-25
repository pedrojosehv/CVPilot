"""
Multi-Template Processor for CVPilot
Extracts and optimizes content from multiple CV templates
"""

import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import JobData, MatchResult
from .logger import LoggerMixin


@dataclass
class TemplateContent:
    """Represents extracted content from a CV template"""
    template_name: str
    template_path: Path
    bullet_points: List[str]
    education: List[str]
    certifications: List[str]
    skills: List[str]
    software: List[str]
    experience_titles: List[str]
    summary: str
    total_length: int


@dataclass
class ContentMatch:
    """Represents a content match with relevance score"""
    content: str
    source_template: str
    relevance_score: float
    content_type: str  # 'bullet', 'education', 'certification', 'skill', 'software'


class MultiTemplateProcessor(LoggerMixin):
    """Processes multiple CV templates to extract optimal content"""
    
    def __init__(self, templates_path: Path):
        super().__init__()
        self.templates_path = templates_path
        self.template_pool: Dict[str, TemplateContent] = {}
        self.logger.info(f"Initializing MultiTemplateProcessor with path: {templates_path}")
    
    def load_template_pool(self) -> Dict[str, TemplateContent]:
        """Load all CV templates and extract their content"""
        self.logger.info("Loading template pool...")
        
        # Find all DOCX files in templates directory
        docx_files = list(self.templates_path.glob("*.docx"))
        self.logger.info(f"Found {len(docx_files)} DOCX templates")
        
        for docx_file in docx_files:
            try:
                template_content = self._extract_template_content(docx_file)
                self.template_pool[docx_file.stem] = template_content
                self.logger.info(f"Loaded template: {docx_file.stem}")
            except Exception as e:
                self.logger.warning(f"Failed to load template {docx_file.name}: {e}")
        
        self.logger.info(f"Successfully loaded {len(self.template_pool)} templates")
        return self.template_pool
    
    def _extract_template_content(self, template_path: Path) -> TemplateContent:
        """Extract content from a single DOCX template"""
        doc = Document(template_path)
        
        bullet_points = []
        education = []
        certifications = []
        skills = []
        software = []
        experience_titles = []
        summary = ""
        total_length = 0
        
        # Extract text and analyze structure
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            total_length += len(text)
            
            # Extract bullet points (lines starting with •, -, *, etc.)
            if re.match(r'^[•\-\*]\s+', text):
                bullet_points.append(text)
            
            # Extract education (look for degree patterns)
            elif any(degree in text.lower() for degree in ['bachelor', 'master', 'phd', 'degree', 'university', 'college']):
                education.append(text)
            
            # Extract certifications (look for cert patterns)
            elif any(cert in text.lower() for cert in ['certification', 'certified', 'certificate', 'pmp', 'csm', 'cspo']):
                certifications.append(text)
            
            # Extract experience titles (job titles in experience section)
            elif self._is_experience_title(text):
                experience_titles.append(text)
            
            # Extract summary (usually first paragraph after title)
            elif len(summary) < 500 and not any(section in text.lower() for section in ['skills', 'experience', 'education']):
                if not summary:
                    summary = text
                else:
                    summary += " " + text
        
        # Extract skills and software from specific sections
        skills, software = self._extract_skills_software(doc)
        
        return TemplateContent(
            template_name=template_path.stem,
            template_path=template_path,
            bullet_points=bullet_points,
            education=education,
            certifications=certifications,
            skills=skills,
            software=software,
            experience_titles=experience_titles,
            summary=summary,
            total_length=total_length
        )
    
    def _is_experience_title(self, text: str) -> bool:
        """Check if text looks like an experience job title"""
        # Look for patterns like "Senior Product Manager", "Lead Developer", etc.
        title_patterns = [
            r'\b(Senior|Lead|Principal|Staff|Junior|Associate)\s+\w+',
            r'\b(Manager|Director|Head|VP|CTO|CEO)\b',
            r'\b(Developer|Engineer|Analyst|Specialist|Consultant)\b'
        ]
        
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in title_patterns)
    
    def _extract_skills_software(self, doc: Document) -> Tuple[List[str], List[str]]:
        """Extract skills and software from document sections"""
        skills = []
        software = []
        
        in_skills_section = False
        in_software_section = False
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Detect sections
            if 'skills' in text.lower():
                in_skills_section = True
                in_software_section = False
                continue
            elif 'software' in text.lower() or 'tools' in text.lower():
                in_software_section = True
                in_skills_section = False
                continue
            elif any(section in text.lower() for section in ['experience', 'education', 'summary']):
                in_skills_section = False
                in_software_section = False
                continue
            
            # Extract content based on section
            if in_skills_section and text:
                skills.extend([s.strip() for s in text.split(';') if s.strip()])
            elif in_software_section and text:
                software.extend([s.strip() for s in text.split(';') if s.strip()])
        
        return skills, software
    
    def select_optimal_content(self, job_data: JobData, match_result: MatchResult, target_length: int) -> Dict[str, Any]:
        """Select optimal content from template pool based on job requirements"""
        self.logger.info(f"Selecting optimal content for job: {job_data.job_title_original}")
        
        # Calculate relevance scores for all content
        content_matches = []
        
        for template_name, template_content in self.template_pool.items():
            # Score bullet points
            for bullet in template_content.bullet_points:
                score = self._calculate_relevance_score(bullet, job_data, match_result)
                content_matches.append(ContentMatch(
                    content=bullet,
                    source_template=template_name,
                    relevance_score=score,
                    content_type='bullet'
                ))
            
            # Score education
            for edu in template_content.education:
                score = self._calculate_relevance_score(edu, job_data, match_result)
                content_matches.append(ContentMatch(
                    content=edu,
                    source_template=template_name,
                    relevance_score=score,
                    content_type='education'
                ))
            
            # Score certifications
            for cert in template_content.certifications:
                score = self._calculate_relevance_score(cert, job_data, match_result)
                content_matches.append(ContentMatch(
                    content=cert,
                    source_template=template_name,
                    relevance_score=score,
                    content_type='certification'
                ))
        
        # Select top content by type
        optimal_content = {
            'bullet_points': self._select_top_content(content_matches, 'bullet', 5),
            'education': self._select_top_content(content_matches, 'education', 2),
            'certifications': self._select_top_content(content_matches, 'certification', 3),
            'skills': self._get_optimal_skills(job_data, match_result),
            'software': self._get_optimal_software(job_data, match_result)
        }
        
        self.logger.info(f"Selected {len(optimal_content['bullet_points'])} bullet points, "
                        f"{len(optimal_content['education'])} education items, "
                        f"{len(optimal_content['certifications'])} certifications")
        
        return optimal_content
    
    def _calculate_relevance_score(self, content: str, job_data: JobData, match_result: MatchResult) -> float:
        """Calculate relevance score for content based on job requirements"""
        score = 0.0
        content_lower = content.lower()
        
        # Score based on job title match
        job_title_words = job_data.job_title_original.lower().split()
        for word in job_title_words:
            if word in content_lower:
                score += 2.0
        
        # Score based on required skills
        for skill in job_data.skills:
            if skill.lower() in content_lower:
                score += 1.5
        
        # Score based on required software
        for software in job_data.software:
            if software.lower() in content_lower:
                score += 1.0
        
        # Score based on matched skills
        for skill in match_result.matched_skills:
            if skill.lower() in content_lower:
                score += 1.0
        
        # Score based on matched software
        for software in match_result.matched_software:
            if software.lower() in content_lower:
                score += 0.8
        
        # Bonus for healthcare/clinical content if job is healthcare-related
        if any(term in job_data.job_title_original.lower() for term in ['healthcare', 'clinical', 'medical']):
            if any(term in content_lower for term in ['healthcare', 'clinical', 'medical', 'patient', 'regulatory']):
                score += 3.0
        
        return score
    
    def _select_top_content(self, content_matches: List[ContentMatch], content_type: str, max_items: int) -> List[str]:
        """Select top content by relevance score"""
        type_matches = [cm for cm in content_matches if cm.content_type == content_type]
        type_matches.sort(key=lambda x: x.relevance_score, reverse=True)
        
        selected = []
        for match in type_matches[:max_items]:
            if match.relevance_score > 0:  # Only include relevant content
                selected.append(match.content)
        
        return selected
    
    def _get_optimal_skills(self, job_data: JobData, match_result: MatchResult) -> List[str]:
        """Get optimal skills combining job requirements and matched skills"""
        optimal_skills = []
        
        # Start with required skills
        optimal_skills.extend(job_data.skills)
        
        # Add matched skills that aren't already included
        for skill in match_result.matched_skills:
            if skill not in optimal_skills:
                optimal_skills.append(skill)
        
        # Limit to top 15 skills
        return optimal_skills[:15]
    
    def _get_optimal_software(self, job_data: JobData, match_result: MatchResult) -> List[str]:
        """Get optimal software combining job requirements and matched software"""
        optimal_software = []
        
        # Start with required software
        optimal_software.extend(job_data.software)
        
        # Add matched software that isn't already included
        for software in match_result.matched_software:
            if software not in optimal_software:
                optimal_software.append(software)
        
        # Limit to top 10 software
        return optimal_software[:10]
    
    def generate_optimized_cv_content(self, job_data: JobData, match_result: MatchResult, base_template: str) -> Dict[str, Any]:
        """Generate optimized CV content using multi-template approach"""
        self.logger.info(f"Generating optimized CV content for {job_data.job_title_original}")
        
        # Load template pool if not already loaded
        if not self.template_pool:
            self.load_template_pool()
        
        # Get base template content for length reference
        base_content = self.template_pool.get(base_template)
        target_length = base_content.total_length if base_content else 5000
        
        # Select optimal content
        optimal_content = self.select_optimal_content(job_data, match_result, target_length)
        
        # Generate optimized summary
        optimized_summary = self._generate_optimized_summary(job_data, match_result, optimal_content)
        
        return {
            'summary': optimized_summary,
            'bullet_points': optimal_content['bullet_points'],
            'education': optimal_content['education'],
            'certifications': optimal_content['certifications'],
            'skills': optimal_content['skills'],
            'software': optimal_content['software'],
            'target_length': target_length
        }
    
    def _generate_optimized_summary(self, job_data: JobData, match_result: MatchResult, optimal_content: Dict[str, Any]) -> str:
        """Generate optimized professional summary"""
        # Combine best elements from template summaries
        summary_elements = []
        
        for template_name, template_content in self.template_pool.items():
            if template_content.summary:
                summary_elements.append(template_content.summary)
        
        # Create a new summary based on job requirements and template elements
        summary = f"Experienced {job_data.job_title_original} with expertise in "
        
        # Add key skills
        key_skills = optimal_content['skills'][:5]
        summary += ", ".join(key_skills) + ". "
        
        # Add key achievements from bullet points
        if optimal_content['bullet_points']:
            summary += f"Demonstrated success in {optimal_content['bullet_points'][0].lower()}. "
        
        # Add software expertise
        if optimal_content['software']:
            summary += f"Proficient in {', '.join(optimal_content['software'][:3])}. "
        
        # Ensure summary is within reasonable length
        if len(summary) > 500:
            summary = summary[:497] + "..."
        
        return summary
