#!/usr/bin/env python3
"""
Analyze the current template structure to understand CV title placement
"""

import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

import docx

def analyze_current_template():
    """Analyze the template that was auto-selected"""

    # Template that was auto-selected (from the list)
    template_path = "output/Project Manager - General - Excel (advanced), Amazon Seller Central, Helium 10/PedroHerrera_PJM_GEN_MP_2025.docx"

    print(f"🔍 Analyzing auto-selected template: {template_path}")

    if not Path(template_path).exists():
        print(f"❌ Template not found: {template_path}")
        print("\nAvailable templates:")
        for template in Path("output").glob("**/*.docx"):
            print(f"   • {template}")
        return

    try:
        doc = docx.Document(template_path)

        print("\n📋 Template Structure Analysis:")
        print("=" * 60)

        print(f"📄 Total paragraphs: {len(doc.paragraphs)}")

        # Look for CV title
        cv_title_found = False
        cv_title_keywords = ['PRODUCT ANALYST', 'CURRICULUM VITAE', 'CV', 'RESUME', 'CURRÍCULUM VITAE', 'INNOVATION SPECIALIST', 'PROJECT MANAGER']

        print("\n🔍 Looking for CV title in first 10 paragraphs:")
        for i, paragraph in enumerate(doc.paragraphs[:10]):
            text = paragraph.text.strip()
            if text:
                print("2d")

                # Check if it matches CV title keywords
                text_upper = text.upper()
                is_cv_title = any(keyword in text_upper for keyword in cv_title_keywords)
                if is_cv_title:
                    print(f"      🎯 CV TITLE FOUND!")
                    cv_title_found = True
                elif 'PROJECT MANAGER' in text_upper:
                    print(f"      🎯 PROJECT MANAGER TITLE FOUND!")
                    cv_title_found = True

        if not cv_title_found:
            print("\n⚠️ No CV title found with current keywords")
            print("\n💡 First few paragraphs contain:")
            for i, paragraph in enumerate(doc.paragraphs[:5]):
                text = paragraph.text.strip()
                if text:
                    print(f"   {i+1}: '{text}'")

        # Look for professional summary
        print("\n🔍 Looking for Professional Summary:")
        summary_keywords = ['PRODUCT MANAGER WITH', 'PROFESSIONAL SUMMARY', 'RESUMEN PROFESIONAL', 'SUMMARY', 'EXECUTIVE SUMMARY']
        summary_found = False

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                text_upper = text.upper()
                is_summary = any(keyword in text_upper for keyword in summary_keywords)
                if is_summary:
                    print(f"   📝 Found at paragraph {i+1}: '{text[:100]}...'")
                    summary_found = True
                    break

        if not summary_found:
            print("   ⚠️ No professional summary found with current keywords")

        # Look for experience titles
        print("\n🔍 Looking for Experience Titles:")
        experience_titles = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text and len(text) < 100 and ('Manager' in text or 'Analyst' in text):
                print(f"   💼 Paragraph {i+1}: '{text}'")
                experience_titles.append(text)

        print(f"\n📊 Found {len(experience_titles)} potential experience titles")

    except Exception as e:
        print(f"❌ Error analyzing template: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_current_template()
