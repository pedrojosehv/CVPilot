#!/usr/bin/env python3
"""
Script para analizar la estructura completa de perfiles y roles múltiples en bullet_pool.docx
"""

import docx
from pathlib import Path

def analyze_bullet_structure():
    """Analizar la estructura completa de perfiles y roles"""
    
    bullet_pool_path = Path("D:/Work Work/Upwork/CVPilot/templates/bullet_pool.docx")
    
    if not bullet_pool_path.exists():
        print(f"❌ Archivo no encontrado: {bullet_pool_path}")
        return
    
    try:
        # Cargar documento
        doc = docx.Document(str(bullet_pool_path))
        
        print("🎯 ESTRUCTURA COMPLETA DEL BULLET_POOL.DOCX")
        print("=" * 60)
        
        # Analizar tablas (roles específicos)
        print("📊 ROLES ESPECÍFICOS EN GCA (Progresión Interna):")
        print("-" * 50)
        
        for i, table in enumerate(doc.tables):
            if table.rows:
                row = table.rows[0]
                if len(row.cells) >= 2:
                    role = row.cells[0].text.strip()
                    period = row.cells[1].text.strip()
                    print(f"👤 Rol {i+1}: {role}")
                    print(f"📅 Período: {period}")
                    print()
        
        # Analizar párrafos (bullet points por compañía)
        print("🏢 BULLET POINTS POR COMPAÑÍA:")
        print("-" * 50)
        
        current_company = None
        bullet_count = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Detectar encabezados de compañía
            if "—" in text and ("Remote" in text or "," in text):
                if current_company:
                    print(f"   Total bullets para {current_company}: {bullet_count}")
                    print()
                
                current_company = text
                bullet_count = 0
                print(f"🏢 {current_company}")
                print("-" * 40)
            else:
                # Es un bullet point
                bullet_count += 1
                print(f"  {bullet_count}. {text}")
        
        if current_company:
            print(f"   Total bullets para {current_company}: {bullet_count}")
        
        print("\n📈 ANÁLISIS DE LA PROGRESIÓN INTERNA EN GCA:")
        print("-" * 50)
        
        # Extraer roles de las tablas
        gca_roles = []
        for table in doc.tables:
            if table.rows:
                row = table.rows[0]
                if len(row.cells) >= 2:
                    role = row.cells[0].text.strip()
                    period = row.cells[1].text.strip()
                    gca_roles.append((role, period))
        
        print("🔄 Progresión de roles en GCA:")
        for i, (role, period) in enumerate(gca_roles):
            print(f"  {i+1}. {role}")
            print(f"     📅 {period}")
        
        print(f"\n📊 RESUMEN:")
        print(f"   • Total de roles en GCA: {len(gca_roles)}")
        print(f"   • Total de compañías: 4")
        print(f"   • Total de bullet points: {sum(1 for p in doc.paragraphs if p.text.strip() and '—' not in p.text)}")
        print(f"   • Estructura: Perfil Avanzado (roles múltiples) + Perfil Básico (roles simples)")
        
    except Exception as e:
        print(f"❌ Error al leer el archivo: {e}")

if __name__ == "__main__":
    analyze_bullet_structure()
