#!/usr/bin/env python3
"""
Script detallado para analizar la estructura completa del bullet_pool.docx
"""

import docx
from pathlib import Path

def analyze_bullet_pool_detailed():
    """Análisis detallado del bullet_pool.docx"""
    
    bullet_pool_path = Path("D:/Work Work/Upwork/CVPilot/templates/bullet_pool.docx")
    
    if not bullet_pool_path.exists():
        print(f"❌ Archivo no encontrado: {bullet_pool_path}")
        return
    
    try:
        # Cargar documento
        doc = docx.Document(str(bullet_pool_path))
        
        print("🔍 ANÁLISIS DETALLADO DEL BULLET_POOL.DOCX")
        print("=" * 60)
        
        print(f"📄 Total de párrafos en el documento: {len(doc.paragraphs)}")
        print(f"📄 Total de tablas en el documento: {len(doc.tables)}")
        print(f"📄 Total de secciones en el documento: {len(doc.sections)}")
        
        print("\n📋 TODOS LOS PÁRRAFOS (con números de línea):")
        print("-" * 50)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"Línea {i+1:2d}: {text}")
            else:
                print(f"Línea {i+1:2d}: [VACÍO]")
        
        print("\n🔍 BÚSQUEDA DE PATRONES ESPECÍFICOS:")
        print("-" * 50)
        
        # Buscar patrones específicos
        patterns = {
            "Perfil": [],
            "Avanzado": [],
            "Básico": [],
            "GCA": [],
            "Growing Companies": [],
            "Role": [],
            "Analyst": [],
            "Manager": []
        }
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            for pattern in patterns.keys():
                if pattern.lower() in text.lower():
                    patterns[pattern].append((i+1, text))
        
        for pattern, matches in patterns.items():
            if matches:
                print(f"\n🔍 '{pattern}' encontrado en:")
                for line_num, text in matches:
                    print(f"   Línea {line_num}: {text}")
        
        # Analizar estructura de tablas si existen
        if doc.tables:
            print(f"\n📊 ANÁLISIS DE TABLAS:")
            print("-" * 30)
            for i, table in enumerate(doc.tables):
                print(f"Tabla {i+1}: {len(table.rows)} filas, {len(table.columns)} columnas")
                for row_idx, row in enumerate(table.rows):
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        print(f"  Fila {row_idx+1}: {' | '.join(row_text)}")
        
    except Exception as e:
        print(f"❌ Error al leer el archivo: {e}")

if __name__ == "__main__":
    analyze_bullet_pool_detailed()
