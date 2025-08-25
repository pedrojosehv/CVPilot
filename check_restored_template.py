"""
Script para verificar el template restaurado
"""

from docx import Document
import os

def check_restored_template():
    print("🔍 VERIFICANDO TEMPLATE RESTAURADO")
    print("=" * 40)

    template_path = "templates/CV Pedro Herrera.docx"

    if not os.path.exists(template_path):
        print(f"❌ Template no encontrado: {template_path}")
        return

    print(f"✅ Template encontrado: {template_path}")

    try:
        doc = Document(template_path)
        print(f"📊 Total de párrafos: {len(doc.paragraphs)}")

        print("\n📋 CONTENIDO DEL TEMPLATE:")
        print("-" * 30)

        titles_found = []

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            if text:
                print(f"[{i:2d}] {text}")

                # Buscar títulos con fechas
                if any(char in text for char in ['/', '-', 'present', '2020', '2021', '2022', '2023']):
                    # Verificar si parece un título
                    if any(word in text.lower() for word in [
                        'analyst', 'manager', 'specialist', 'consultant', 'developer',
                        'engineer', 'director', 'coordinator'
                    ]):
                        titles_found.append((i, text))
                        print("     🎯 TÍTULO DETECTADO")

        print("\n🎯 TÍTULOS ENCONTRADOS:")
        print("-" * 30)

        for i, (line_num, title) in enumerate(titles_found, 1):
            print(f"{i}. Línea {line_num}: '{title}'")

        print("\n📊 RESUMEN:")
        print(f"✅ Template restaurado exitosamente")
        print(f"✅ Títulos encontrados: {len(titles_found)}")

        if len(titles_found) > 0:
            print("✅ EL SISTEMA DEBERÍA FUNCIONAR AHORA")
        else:
            print("⚠️  Todavía no hay títulos específicos")

    except Exception as e:
        print(f"❌ Error al leer el template: {e}")

if __name__ == "__main__":
    check_restored_template()
